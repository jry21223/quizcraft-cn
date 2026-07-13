#!/usr/bin/env python3
# -*- coding: UTF-8 -*
"""
刷题系统后端 - FastAPI 版本
提供 RESTful API 和 WebSocket 实时推送
"""

import json
import os
import re
import tempfile
import shutil
import asyncio
import hashlib
import hmac
import secrets
import time
from typing import Dict, List, Optional, Any, Set, Tuple
from datetime import datetime, timedelta
from collections import defaultdict
from contextlib import asynccontextmanager, suppress
from pathlib import Path
from zoneinfo import ZoneInfo

from fastapi import Depends, FastAPI, Header, HTTPException, Request, Response, UploadFile, File, Form, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

import db_storage
from scripts.java_bank_workflow import (
    DEFAULT_JAVA_BANK_KEY,
    DEFAULT_JAVA_START_NUMBER,
    build_incremental_bank,
    fill_java_questions_with_deepseek,
    parse_java_markdown_bank,
)


# CORS 配置辅助函数：读取环境变量 CORS_ORIGINS（逗号分隔），默认只允许本地开发地址
def get_cors_origins() -> list:
    raw = os.getenv("CORS_ORIGINS", "").strip()
    if not raw:
        return [
            "http://127.0.0.1:5173",
            "http://localhost:5173",
        ]
    return [origin.strip() for origin in raw.split(",") if origin.strip()]


def get_admin_token() -> str:
    return os.getenv("ADMIN_TOKEN", "").strip()


ADMIN_SESSION_COOKIE_NAME = "quizcraft_admin_session"
DEFAULT_ADMIN_SESSION_TTL_SECONDS = 8 * 60 * 60


def get_admin_session_ttl_seconds() -> int:
    raw = os.getenv("ADMIN_SESSION_TTL_SECONDS", "").strip()
    try:
        ttl = int(raw) if raw else DEFAULT_ADMIN_SESSION_TTL_SECONDS
    except ValueError:
        ttl = DEFAULT_ADMIN_SESSION_TTL_SECONDS
    return min(max(ttl, 5 * 60), 24 * 60 * 60)


def _admin_session_signature(payload: str) -> str:
    signing_key = hashlib.sha256(
        f"quizcraft-admin-session:{get_admin_token()}".encode("utf-8")
    ).digest()
    return hmac.new(signing_key, payload.encode("utf-8"), hashlib.sha256).hexdigest()


def create_admin_session_token(now: Optional[int] = None) -> str:
    if not get_admin_token():
        raise RuntimeError("后台管理 Token 未配置")
    issued_at = int(time.time() if now is None else now)
    payload = f"{issued_at + get_admin_session_ttl_seconds()}.{secrets.token_urlsafe(24)}"
    return f"{payload}.{_admin_session_signature(payload)}"


def is_admin_session_valid(token: Optional[str], now: Optional[int] = None) -> bool:
    if not get_admin_token() or not token:
        return False
    try:
        payload, signature = token.rsplit(".", 1)
        expires_at_text, _nonce = payload.split(".", 1)
        expires_at = int(expires_at_text)
    except (TypeError, ValueError):
        return False

    expected_signature = _admin_session_signature(payload)
    if not secrets.compare_digest(signature, expected_signature):
        return False
    current_time = int(time.time() if now is None else now)
    return expires_at > current_time


def should_use_secure_admin_cookie(request: Request) -> bool:
    forwarded_scheme = request.headers.get("x-forwarded-proto", "").split(",", 1)[0].strip()
    scheme = forwarded_scheme or request.url.scheme
    if scheme == "https":
        return True
    return request.url.hostname not in {"127.0.0.1", "localhost", "testserver"}


def get_disabled_bank_keys() -> Set[str]:
    raw = os.getenv("DISABLED_BANK_KEYS", "").strip()
    if not raw:
        return set()
    return {
        key.strip()
        for key in re.split(r"[,;\s]+", raw)
        if key.strip()
    }


def should_sync_local_banks_to_db() -> bool:
    raw = os.getenv("QUIZCRAFT_SYNC_LOCAL_BANKS_TO_DB", "").strip().lower()
    return raw in {"1", "true", "yes", "on"}


def is_bank_enabled(bank_key: str) -> bool:
    return bank_key not in get_disabled_bank_keys()


def require_enabled_bank(bank_key: str):
    if bank_key not in QUESTION_BANKS or not is_bank_enabled(bank_key):
        raise HTTPException(status_code=404, detail="题库不存在")


def is_admin_token_valid(token: Optional[str]) -> bool:
    expected = get_admin_token()
    return bool(expected and token and secrets.compare_digest(token, expected))


async def require_admin_token(
    request: Request,
    x_admin_token: Optional[str] = Header(None),
):
    if not get_admin_token():
        raise HTTPException(status_code=503, detail="后台管理 Token 未配置")
    session_token = request.cookies.get(ADMIN_SESSION_COOKIE_NAME)
    if not (
        is_admin_session_valid(session_token)
        or is_admin_token_valid(x_admin_token)
    ):
        raise HTTPException(status_code=403, detail="后台管理 Token 无效")

# 导入 LLM 服务
try:
    from llm_service import LLMService, LLMConfig
except ImportError:
    LLMService = None
    LLMConfig = None


# WebSocket 连接管理
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}
    
    async def connect(self, client_id: str, websocket: WebSocket):
        await websocket.accept()
        self.active_connections[client_id] = websocket
    
    def disconnect(self, client_id: str):
        if client_id in self.active_connections:
            del self.active_connections[client_id]
    
    async def send_progress(self, client_id: str, current: int, total: int, message: str = ""):
        if client_id in self.active_connections:
            await self.active_connections[client_id].send_json({
                "type": "progress",
                "current": current,
                "total": total,
                "percentage": round(current / total * 100, 1),
                "message": message
            })
    
    async def send_complete(self, client_id: str, questions: List[Dict]):
        if client_id in self.active_connections:
            await self.active_connections[client_id].send_json({
                "type": "complete",
                "questions": questions
            })
    
    async def send_error(self, client_id: str, error: str):
        if client_id in self.active_connections:
            await self.active_connections[client_id].send_json({
                "type": "error",
                "error": error
            })

manager = ConnectionManager()


# ============== 数据模型 ==============

class PracticeSettings(BaseModel):
    mode: str = "random"
    params: Dict[str, Any] = {}


class StartPracticeRequest(BaseModel):
    bank: str
    mode: str
    params: Dict[str, Any]


class SubmitAnswerRequest(BaseModel):
    bank: str
    question_id: str
    answer: Any
    user_id: Optional[str] = None


class FeedbackRequest(BaseModel):
    question_index: int
    suggestion: str
    source_page: str = "quiz"
    question_bank: Optional[str] = None
    question_id: Optional[str] = None
    question_content: Optional[str] = None


class FeedbackStatusRequest(BaseModel):
    status: str
    resolution_note: Optional[str] = None


class UserRequest(BaseModel):
    name: Optional[str] = None


class FoodWheelRequest(BaseModel):
    items: List[str]
    user_id: str


class AnalysisConfig(BaseModel):
    provider: str = "deepseek"
    apiKey: str = ""
    apiUrl: Optional[str] = None
    model: Optional[str] = None
    # 可选：支持多 API 配置并发（每个配置可使用不同 provider）
    apiConfigs: Optional[List[Dict[str, Any]]] = None


class AnalyzeRequest(BaseModel):
    questions: List[Dict]
    config: AnalysisConfig


class ExportRequest(BaseModel):
    questions: List[Dict]
    name: str


class SaveBankRequest(BaseModel):
    name: str
    questions: List[Dict]
    key: Optional[str] = None
    color: Optional[str] = None
    overwrite: bool = False


# ============== 全局状态 ==============

QUESTION_BANKS: Dict[str, Dict] = {}
QUESTION_CACHE: Dict[str, List[Dict[str, Any]]] = {}
QUESTION_INDEX: Dict[str, Dict[str, Dict[str, Any]]] = {}
USER_STATS: Dict[str, Dict] = defaultdict(lambda: {
    "name": "",
    "correct": 0,
    "total": 0,
    "practice_history": []
})
NAME_TO_ID: Dict[str, str] = {}
NEXT_USER_ID = 1
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TIKU_DIR = os.path.join(BASE_DIR, "tiku")
EXPORT_DIR = os.path.join(BASE_DIR, "exports")
RANK_FILE = os.path.join(BASE_DIR, "rankings_v2.json")
QUESTION_STATS_FILE = os.path.join(BASE_DIR, "question_stats.json")
FEEDBACK_FILE = os.path.join(BASE_DIR, "feedbacks.json")


def make_generated_user_id(sequence: int) -> str:
    return db_storage.generated_user_id(sequence)


def allocate_generated_user_id() -> str:
    global NEXT_USER_ID
    while True:
        user_id = make_generated_user_id(NEXT_USER_ID)
        NEXT_USER_ID += 1
        if user_id not in USER_STATS:
            return user_id


FOOD_WHEEL_FILE = os.path.join(BASE_DIR, "food_wheel_items.json")
API_CONFIG_CACHE: Dict[str, Tuple[float, List["LLMConfig"]]] = {}
API_CONFIG_CACHE_TTL = 30 * 60  # 30 分钟
QUESTION_GLOBAL_STATS: Dict[str, Dict[str, Dict[str, float]]] = defaultdict(dict)
BANK_COLOR_PALETTE = [
    "#1976d2",
    "#2e7d32",
    "#c62828",
    "#6a1b9a",
    "#ef6c00",
    "#00838f",
    "#5d4037",
    "#7b1fa2",
]
DEFAULT_FOOD_WHEEL_ITEMS = [
    "板面",
    "香扒饭",
    "摇滚炒鸡",
    "盖浇饭",
    "烤肉拌饭",
    "麻辣烫",
    "麦当劳",
]

JUDGE_TRUE_VALUES = {
    "true", "t", "1", "yes", "y", "right",
    "对", "正确", "是", "√",
}
JUDGE_FALSE_VALUES = {
    "false", "f", "0", "no", "n", "wrong",
    "错", "错误", "否", "×",
}


# ============== 题库加载 ==============

def normalize_bank_color(color: Optional[str], fallback_key: str = "") -> str:
    if isinstance(color, str) and re.fullmatch(r"#[0-9a-fA-F]{6}", color.strip()):
        return color.strip()
    if not fallback_key:
        return BANK_COLOR_PALETTE[0]
    digest = hashlib.md5(fallback_key.encode("utf-8")).hexdigest()
    return BANK_COLOR_PALETTE[int(digest[:8], 16) % len(BANK_COLOR_PALETTE)]


def sanitize_bank_key(raw: Optional[str]) -> str:
    text = str(raw or "").strip().lower()
    safe = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", text).strip("_")
    return safe or f"bank_{int(time.time())}"


def register_question_bank(
    key: str,
    name: str,
    color: Optional[str],
    file_path: str,
    data: Any,
    files: Optional[List[str]] = None,
):
    QUESTION_BANKS[key] = {
        "name": name,
        "files": files or [os.path.basename(file_path)],
        "color": normalize_bank_color(color, key),
        "file": file_path,
        "data": data,
    }


def refresh_question_cache(bank_key: Optional[str] = None):
    """预解析题库并建立 question_id 索引，避免每次答题都扫描整套题库。"""
    target_keys = [bank_key] if bank_key else list(QUESTION_BANKS.keys())
    if bank_key is None:
        QUESTION_CACHE.clear()
        QUESTION_INDEX.clear()

    for key in target_keys:
        bank = QUESTION_BANKS.get(key)
        if not bank:
            QUESTION_CACHE.pop(key, None)
            QUESTION_INDEX.pop(key, None)
            continue
        questions = parse_question_bank(bank["data"], key)
        QUESTION_CACHE[key] = questions
        QUESTION_INDEX[key] = {
            str(question.get("id")): question
            for question in questions
            if question.get("id") is not None
        }


def get_bank_questions(bank_key: str) -> List[Dict[str, Any]]:
    if bank_key not in QUESTION_CACHE:
        refresh_question_cache(bank_key)
    return QUESTION_CACHE.get(bank_key, [])


def get_bank_question(bank_key: str, question_id: str) -> Optional[Dict[str, Any]]:
    if bank_key not in QUESTION_INDEX:
        refresh_question_cache(bank_key)
    return QUESTION_INDEX.get(bank_key, {}).get(str(question_id))


def update_cached_question_stats(bank_key: str, question_id: str, stat: Dict[str, Any]):
    question = get_bank_question(bank_key, question_id)
    if not question:
        return
    question["stats"] = {
        "total": int(stat.get("total", 0)),
        "correct": int(stat.get("correct", 0)),
        "rate": float(stat.get("rate", 0)),
    }


def load_bank_from_file(
    key: str,
    file_path: str,
    name: Optional[str] = None,
    color: Optional[str] = None,
    files: Optional[List[str]] = None,
) -> bool:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if not isinstance(data, dict):
            raise ValueError("题库文件顶层必须是 JSON 对象")

        meta = data.get("meta", {}) if isinstance(data, dict) else {}
        register_question_bank(
            key=key,
            name=str(name or meta.get("name") or key),
            color=color or meta.get("color"),
            file_path=file_path,
            data=data,
            files=files,
        )
        print(f"✓ 加载题库: {QUESTION_BANKS[key]['name']} ({key}) <- {file_path}")
        return True
    except Exception as e:
        print(f"✗ 加载失败 {key}: {e}")
        return False


def build_bank_summary(key: str, bank: Dict[str, Any]) -> Dict[str, Any]:
    data = bank["data"]
    chapters = []

    parsed_questions = get_bank_questions(key)
    seen = set()
    for q in parsed_questions:
        ch_name = q.get("chapter")
        ch_id = q.get("chapter_id")
        if not ch_name:
            continue
        if not ch_id:
            ch_id = str(ch_name)
        if ch_id not in seen:
            chapters.append({"id": ch_id, "name": ch_name})
            seen.add(ch_id)

    if not chapters:
        chapters = [{"id": "ch01", "name": "默认章节"}]

    total = 0
    if isinstance(data, dict) and "questions" in data:
        total = len(parsed_questions)
    elif isinstance(data, dict) and "meta" in data:
        total = data["meta"].get("total", 0)
    elif isinstance(data, dict):
        total = sum(
            len(items)
            for types in data.values()
            if isinstance(types, dict)
            for items in types.values()
            if isinstance(items, list)
        )

    return {
        "key": key,
        "name": bank["name"],
        "color": bank["color"],
        "total": total,
        "chapters": chapters,
    }


def build_standard_bank_data(name: str, questions: List[Dict], color: Optional[str] = None) -> Dict[str, Any]:
    normalized_questions: List[Dict[str, Any]] = []
    chapter_to_id: Dict[str, str] = {}

    for idx, raw in enumerate(questions, start=1):
        if not isinstance(raw, dict):
            continue

        content = str(
            raw.get("content")
            or raw.get("question")
            or raw.get("title")
            or ""
        ).strip()
        if not content:
            raise HTTPException(status_code=400, detail=f"第 {idx} 题题干不能为空")

        chapter_name = str(
            raw.get("chapter")
            or raw.get("chapterName")
            or raw.get("section")
            or raw.get("group")
            or raw.get("章节")
            or raw.get("组别")
            or "默认章节"
        ).strip() or "默认章节"
        chapter_id = str(raw.get("chapter_id") or "").strip()
        if not chapter_id:
            chapter_id = chapter_to_id.setdefault(chapter_name, f"ch{len(chapter_to_id) + 1:02d}")

        q_type = _normalize_q_type(raw.get("type"), _answer_to_text(raw.get("answer"), "single"))
        analysis = str(raw.get("analysis") or "").strip()

        question: Dict[str, Any] = {
            "id": str(raw.get("id") or f"q{idx:04d}"),
            "number": str(raw.get("number") or idx),
            "type": q_type,
            "content": content,
            "analysis": analysis,
            "chapter": chapter_name,
            "chapter_id": chapter_id,
            "stats": {
                "total": 0,
                "correct": 0,
                "rate": 0,
            },
        }

        if q_type == "blank":
            normalized_answer = normalize_blank_answer_value(raw.get("answer"))
            if not normalized_answer:
                raise HTTPException(status_code=400, detail=f"第 {idx} 题填空题答案不能为空")
            question["answer"] = normalized_answer
            question["options"] = None
            normalized_questions.append(question)
            continue

        if q_type == "judge":
            normalized_answer = normalize_judge_answer(raw.get("answer"))
            if normalized_answer is None:
                raise HTTPException(status_code=400, detail=f"第 {idx} 题判断题答案必须是“对/错”")
            question["answer"] = normalized_answer
            question["options"] = None
            normalized_questions.append(question)
            continue

        options = [opt for opt in _normalize_options(raw.get("options")) if opt]
        if len(options) < 2:
            raise HTTPException(status_code=400, detail=f"第 {idx} 题至少需要 2 个有效选项")

        normalized_type, normalized_answer = normalize_choice_answer(
            raw.get("answer"),
            q_type,
            options,
            analysis,
            content,
        )
        if normalized_answer is None:
            raise HTTPException(status_code=400, detail=f"第 {idx} 题答案格式无效，请检查正确答案")

        question["type"] = normalized_type
        question["options"] = options
        question["answer"] = normalized_answer
        normalized_questions.append(question)

    if not normalized_questions:
        raise HTTPException(status_code=400, detail="题库中没有可保存的题目")

    meta = {
        "name": name,
        "version": "1.0.0",
        "created_at": datetime.now().isoformat(),
        "total": len(normalized_questions),
        "chapter_count": len(chapter_to_id) or 1,
    }
    normalized_color = normalize_bank_color(color, name)
    if normalized_color:
        meta["color"] = normalized_color

    return {
        "meta": meta,
        "questions": normalized_questions,
    }

def register_db_question_banks() -> int:
    loaded_from_db = 0
    for key, db_bank in db_storage.load_question_banks().items():
        metadata = dict(db_bank.get("metadata") or {})
        data = {
            "meta": metadata,
            "questions": db_bank.get("questions") or [],
        }
        register_question_bank(
            key=key,
            name=str(db_bank.get("name") or metadata.get("name") or key),
            color=str(db_bank.get("color") or metadata.get("color") or ""),
            file_path=f"postgresql:{key}",
            data=data,
            files=[],
        )
        loaded_from_db += 1
    return loaded_from_db


def load_question_banks():
    """加载所有题库：生产 PostgreSQL 优先，本地 tiku JSON 仅作兜底或一次性迁移来源。"""
    QUESTION_BANKS.clear()
    QUESTION_CACHE.clear()
    QUESTION_INDEX.clear()
    os.makedirs(TIKU_DIR, exist_ok=True)

    if db_runtime_enabled() and not should_sync_local_banks_to_db():
        try:
            loaded_from_db = register_db_question_banks()
            if loaded_from_db:
                print(f"✓ 从 PostgreSQL 加载 {loaded_from_db} 个题库")
                refresh_question_cache()
                return
        except Exception as e:
            print(f"从 PostgreSQL 加载题库失败，将回退到本地 JSON: {e}")

    banks = {
        "sixiu": {
            "name": "思想道德与法治",
            "files": ["sixiu_with_stats.json", "sixiu.json"],
            "color": "#2e7d32"
        },
        "xigai": {
            "name": "习概",
            "files": ["xigai.json"],
            "color": "#1976d2"
        },
        "history": {
            "name": "近代史",
            "files": ["history.json", "chapters.json"],
            "color": "#c62828"
        },
        "software_engineering_process_tests": {
            "name": "软件工程过程性测试",
            "files": ["generated/software_engineering_process_tests.json"],
            "color": "#1565c0"
        }
    }

    loaded_files: Set[str] = set()
    for key, config in banks.items():
        file_path = None
        candidates = config.get("files", [])
        # 兼容：优先项目根目录，其次 tiku 目录
        for candidate in candidates:
            candidate_paths = [
                os.path.join(BASE_DIR, candidate),
                os.path.join(TIKU_DIR, candidate),
            ]
            for candidate_path in candidate_paths:
                if os.path.exists(candidate_path):
                    file_path = candidate_path
                    break
            if file_path:
                break

        if file_path and os.path.exists(file_path):
            if load_bank_from_file(
                key=key,
                file_path=file_path,
                name=config["name"],
                color=config["color"],
                files=config.get("files"),
            ):
                loaded_files.add(os.path.abspath(file_path))

    for filename in sorted(os.listdir(TIKU_DIR)):
        if not filename.endswith(".json"):
            continue
        file_path = os.path.join(TIKU_DIR, filename)
        abs_path = os.path.abspath(file_path)
        key = os.path.splitext(filename)[0]
        if abs_path in loaded_files or key in QUESTION_BANKS:
            continue
        if load_bank_from_file(key=key, file_path=file_path):
            loaded_files.add(abs_path)

    if db_runtime_enabled():
        try:
            loaded_from_db = 0
            for key, db_bank in db_storage.load_question_banks().items():
                if key in QUESTION_BANKS:
                    continue
                metadata = dict(db_bank.get("metadata") or {})
                data = {
                    "meta": metadata,
                    "questions": db_bank.get("questions") or [],
                }
                register_question_bank(
                    key=key,
                    name=str(db_bank.get("name") or metadata.get("name") or key),
                    color=str(db_bank.get("color") or metadata.get("color") or ""),
                    file_path=f"postgresql:{key}",
                    data=data,
                    files=[],
                )
                loaded_from_db += 1
            if loaded_from_db:
                print(f"✓ 从 PostgreSQL 加载 {loaded_from_db} 个题库")
        except Exception as e:
            print(f"从 PostgreSQL 加载题库失败: {e}")

    refresh_question_cache()


def db_runtime_enabled() -> bool:
    return db_storage.is_available()


def initialize_database_if_configured():
    if not db_storage.is_enabled():
        print("ℹ️ 未配置 DATABASE_URL，使用本地 JSON 运行时兜底")
        return
    if not db_storage.is_available():
        raise RuntimeError("DATABASE_URL 已配置，但 psycopg 不可用，请安装 requirements.txt")
    db_storage.init_schema()
    print("✓ PostgreSQL 存储已初始化")


def sync_question_bank_to_db(key: str, bank: Dict[str, Any]) -> bool:
    if not db_runtime_enabled():
        return False
    source_file = str(bank.get("file") or "")
    if source_file.startswith("postgresql:"):
        return False
    data = bank["data"]
    metadata = data.get("meta", {}) if isinstance(data, dict) else {}
    questions = parse_question_bank(data, key)
    db_storage.upsert_question_bank(
        bank_key=key,
        name=bank["name"],
        color=bank["color"],
        source_file=os.path.relpath(source_file, BASE_DIR),
        metadata=metadata,
        questions=questions,
    )
    return True


def sync_question_banks_to_db():
    if not db_runtime_enabled():
        return
    if not should_sync_local_banks_to_db():
        print("ℹ️ 已跳过本地题库同步到 PostgreSQL；如需一次性迁移请设置 QUIZCRAFT_SYNC_LOCAL_BANKS_TO_DB=1")
        return
    synced = 0
    for key, bank in QUESTION_BANKS.items():
        if sync_question_bank_to_db(key, bank):
            synced += 1
    print(f"✓ 已同步 {synced} 个本地题库到 PostgreSQL")


def save_rankings():
    """保存排行榜"""
    if db_runtime_enabled():
        try:
            db_storage.save_user_snapshot(dict(USER_STATS))
        except Exception as e:
            print(f"保存排行榜到 PostgreSQL 失败: {e}")
        return
    try:
        data = {
            "users": dict(USER_STATS),
            "name_to_id": NAME_TO_ID,
        }
        with open(RANK_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"保存排行榜失败: {e}")


def load_rankings():
    """加载排行榜"""
    global NEXT_USER_ID
    if db_runtime_enabled():
        try:
            users, name_to_id, next_user_id = db_storage.load_runtime_state()
            USER_STATS.clear()
            USER_STATS.update(users)
            NAME_TO_ID.clear()
            NAME_TO_ID.update(name_to_id)
            NEXT_USER_ID = next_user_id
        except Exception as e:
            print(f"从 PostgreSQL 加载排行榜失败: {e}")
        return
    if not os.path.exists(RANK_FILE):
        return
    try:
        with open(RANK_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        users = data.get("users", {})
        name_to_id = data.get("name_to_id", {})
        USER_STATS.clear()
        USER_STATS.update(users)
        NAME_TO_ID.clear()
        NAME_TO_ID.update(name_to_id)
        if users:
            sequence_ids = [
                sequence
                for uid in users.keys()
                if (sequence := db_storage.user_id_sequence_number(uid)) is not None
            ]
            NEXT_USER_ID = max(sequence_ids + [0]) + 1
    except Exception as e:
        print(f"加载排行榜失败: {e}")


def save_question_stats():
    """保存全站题目统计"""
    if db_runtime_enabled():
        return
    try:
        with open(QUESTION_STATS_FILE, 'w', encoding='utf-8') as f:
            json.dump(dict(QUESTION_GLOBAL_STATS), f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"保存题目统计失败: {e}")


def _normalize_food_wheel_items(values) -> List[str]:
    if not isinstance(values, list):
        return []

    normalized: List[str] = []
    seen = set()
    for value in values:
        text = str(value or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        normalized.append(text)
    return normalized


def _load_local_wheels() -> List[Dict[str, Any]]:
    if not os.path.exists(FOOD_WHEEL_FILE):
        return []

    try:
        with open(FOOD_WHEEL_FILE, "r", encoding="utf-8") as f:
            payload = json.load(f)
    except Exception as e:
        print(f"加载美食转盘本地文件失败: {e}")
        return []

    raw_records: List[Any] = []
    if isinstance(payload, dict):
        records = payload.get("wheels")
        if isinstance(records, list):
            raw_records = records
        elif isinstance(payload.get("items"), list):
            legacy = payload.get("items")
            raw_records = [
                {
                    "id": 1,
                    "owner_user_id": "legacy",
                    "owner_name": "legacy",
                    "items": legacy,
                    "is_public": True,
                }
            ]
    elif isinstance(payload, list):
        raw_records = payload

    normalized_records: List[Dict[str, Any]] = []
    now = datetime.now().isoformat()
    for raw in raw_records:
        if not isinstance(raw, dict):
            continue

        owner_user_id = str(raw.get("owner_user_id") or "legacy").strip()
        items = _normalize_food_wheel_items(raw.get("items"))
        is_public = bool(raw.get("is_public", True))
        if not is_public:
            continue
        try:
            wheel_id = int(raw.get("id"))
        except Exception:
            wheel_id = 0

        normalized_records.append(
            {
                "id": wheel_id,
                "owner_user_id": owner_user_id,
                "owner_name": str(raw.get("owner_name") or owner_user_id).strip() or owner_user_id,
                "items": items,
                "is_public": True,
                "created_at": str(raw.get("created_at") or now),
                "updated_at": str(raw.get("updated_at") or now),
            }
        )

    normalized_records.sort(key=lambda item: item.get("updated_at", now), reverse=True)
    return normalized_records


def save_food_wheel_items(user_id: str, values: List[str]) -> Dict[str, Any]:
    normalized = _normalize_food_wheel_items(values)
    user_id = str(user_id or "").strip()
    normalized_name = user_id or "legacy"
    now = datetime.now().isoformat()

    if db_runtime_enabled():
        return db_storage.upsert_food_wheel_items(user_id, normalized)

    records = _load_local_wheels()
    next_id = 1
    updated = False

    for i, record in enumerate(records):
        if str(record.get("owner_user_id") or "") == user_id:
            record["items"] = normalized
            record["owner_name"] = normalized_name
            record["is_public"] = True
            record["updated_at"] = now
            if not record.get("created_at"):
                record["created_at"] = now
            records[i] = record
            updated = True
            break
    else:
        max_id = 0
        for record in records:
            try:
                max_id = max(max_id, int(record.get("id", 0)))
            except Exception:
                continue
        next_id = max_id + 1
        records.append(
            {
                "id": next_id,
                "owner_user_id": user_id,
                "owner_name": normalized_name,
                "items": normalized,
                "is_public": True,
                "created_at": now,
                "updated_at": now,
            }
        )

    records.sort(key=lambda item: str(item.get("updated_at") or now), reverse=True)
    payload = {"wheels": records}

    with open(FOOD_WHEEL_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    if updated:
        for record in records:
            if str(record.get("owner_user_id") or "") == user_id:
                return record

    return {
        "id": next_id,
        "owner_user_id": user_id,
        "owner_name": normalized_name,
        "items": normalized,
        "is_public": True,
        "created_at": now,
        "updated_at": now,
    }


def load_question_stats():
    """加载全站题目统计"""
    if db_runtime_enabled():
        try:
            QUESTION_GLOBAL_STATS.clear()
            QUESTION_GLOBAL_STATS.update(db_storage.load_question_stats())
        except Exception as e:
            print(f"从 PostgreSQL 加载题目统计失败: {e}")
        return
    if not os.path.exists(QUESTION_STATS_FILE):
        return
    try:
        with open(QUESTION_STATS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        QUESTION_GLOBAL_STATS.clear()
        for bank, stats in data.items():
            if isinstance(stats, dict):
                QUESTION_GLOBAL_STATS[bank] = stats
    except Exception as e:
        print(f"加载题目统计失败: {e}")


def load_food_wheel_items():
    """加载美食转盘（仅公开列表）"""
    if db_runtime_enabled():
        try:
            return db_storage.list_food_wheel_items()
        except Exception as e:
            print(f"从 PostgreSQL 加载美食转盘失败: {e}")
            return []

    return _load_local_wheels()


def _normalize_feedback_suggestion(value: str) -> str:
    text = (value or "").strip()
    return text


def _normalize_feedback_bank(value: Optional[str]) -> Optional[str]:
    text = (value or "").strip()
    return text or None


def _normalize_feedback_status(value: Optional[str]) -> str:
    text = (value or "pending").strip().lower()
    return text if text in {"pending", "resolved", "archived"} else "pending"


def _feedback_today_range():
    tz = ZoneInfo("Asia/Shanghai")
    start = datetime.now(tz).replace(hour=0, minute=0, second=0, microsecond=0)
    end = start + timedelta(days=1)
    return start, end


def _parse_feedback_time(value: Any) -> Optional[datetime]:
    if not value:
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(text)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=ZoneInfo("Asia/Shanghai"))
    return parsed


def _serialize_feedback_item(item: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "feedback_id": int(item.get("feedback_id") or 0),
        "question_index": int(item.get("question_index") or 0),
        "question_id": (item.get("question_id") or None),
        "question_content": (item.get("question_content") or None),
        "question_bank": _normalize_feedback_bank(item.get("question_bank")),
        "suggestion": str(item.get("suggestion") or ""),
        "user_id": (item.get("user_id") or None),
        "source_page": str(item.get("source_page") or "quiz") or "quiz",
        "created_at": str(item.get("created_at") or ""),
        "status": _normalize_feedback_status(item.get("status")),
        "resolved_at": str(item.get("resolved_at") or "") or None,
        "resolution_note": str(item.get("resolution_note") or ""),
    }


def load_feedback_dashboard_fallback(
    pending_limit: int = 100,
    resolved_limit: int = 100,
    archived_limit: int = 100,
) -> Dict[str, Any]:
    if not os.path.exists(FEEDBACK_FILE):
        return {
            "summary": {
                "today_total": 0,
                "pending_total": 0,
                "resolved_total": 0,
                "archived_total": 0,
            },
            "pending_items": [],
            "resolved_items": [],
            "archived_items": [],
        }

    try:
        with open(FEEDBACK_FILE, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, json.JSONDecodeError):
        raw = []

    start, end = _feedback_today_range()
    items = [
        _serialize_feedback_item(item)
        for item in raw
        if isinstance(item, dict)
    ]
    today_total = 0
    for item in items:
        created_at = _parse_feedback_time(item.get("created_at"))
        if created_at and start <= created_at < end:
            today_total += 1

    pending_items = [
        item for item in items if item["status"] == "pending"
    ]
    resolved_items = [
        item for item in items if item["status"] == "resolved"
    ]
    archived_items = [
        item for item in items if item["status"] == "archived"
    ]
    pending_items.sort(key=lambda item: (item.get("created_at") or "", item["feedback_id"]), reverse=True)
    resolved_items.sort(
        key=lambda item: (
            item.get("resolved_at") or "",
            item.get("created_at") or "",
            item["feedback_id"],
        ),
        reverse=True,
    )
    archived_items.sort(key=lambda item: (item.get("created_at") or "", item["feedback_id"]), reverse=True)

    return {
        "summary": {
            "today_total": today_total,
            "pending_total": len(pending_items),
            "resolved_total": len(resolved_items),
            "archived_total": len(archived_items),
        },
        "pending_items": pending_items[:pending_limit],
        "resolved_items": resolved_items[:resolved_limit],
        "archived_items": archived_items[:archived_limit],
    }


def update_feedback_status_fallback(
    feedback_id: int,
    status: str,
    resolution_note: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    if not os.path.exists(FEEDBACK_FILE):
        return None
    try:
        with open(FEEDBACK_FILE, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(raw, list):
        return None

    normalized_status = _normalize_feedback_status(status)
    note = (resolution_note or "").strip()[:1000]
    updated_item = None
    for item in raw:
        if not isinstance(item, dict):
            continue
        if int(item.get("feedback_id") or 0) != int(feedback_id):
            continue
        item["status"] = normalized_status
        item["resolution_note"] = note
        item["resolved_at"] = datetime.now(ZoneInfo("Asia/Shanghai")).isoformat() if normalized_status == "resolved" else None
        updated_item = _serialize_feedback_item(item)
        break

    if updated_item is None:
        return None
    with open(FEEDBACK_FILE, "w", encoding="utf-8") as f:
        json.dump(raw, f, ensure_ascii=False, indent=2)
    return updated_item


def save_feedback_fallback(
    question_index: int,
    suggestion: str,
    question_bank: Optional[str] = None,
    question_id: Optional[str] = None,
    question_content: Optional[str] = None,
    source_page: str = "quiz",
) -> Dict[str, Any]:
    normalized_suggestion = _normalize_feedback_suggestion(suggestion)
    if not normalized_suggestion:
        raise ValueError("suggestion is required")
    if len(normalized_suggestion) > 2000:
        raise ValueError("suggestion is too long")

    next_id = 1
    payload = []
    if os.path.exists(FEEDBACK_FILE):
        try:
            with open(FEEDBACK_FILE, "r", encoding="utf-8") as f:
                existing = json.load(f)
                if isinstance(existing, list):
                    payload = existing
        except (OSError, json.JSONDecodeError):
            payload = []
    payload = [item for item in payload if isinstance(item, dict)]
    if payload:
        tail_id = payload[-1].get("feedback_id")
        try:
            next_id = int(tail_id) + 1
        except (TypeError, ValueError):
            next_id = 1

    now = datetime.now().isoformat()
    record = {
        "feedback_id": next_id,
        "question_index": int(question_index),
        "question_id": (question_id or "").strip() or None,
        "question_content": (question_content or "").strip() or None,
        "suggestion": normalized_suggestion,
        "user_id": None,
        "question_bank": _normalize_feedback_bank(question_bank),
        "source_page": (source_page or "quiz").strip() or "quiz",
        "created_at": now,
        "status": "pending",
        "resolved_at": None,
        "resolution_note": "",
    }
    payload.append(record)
    with open(FEEDBACK_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return record


def apply_global_question_stats(bank_key: str, questions: List[Dict]) -> List[Dict]:
    """将全站统计覆盖到题目列表"""
    bank_stats = QUESTION_GLOBAL_STATS.get(bank_key, {})
    if not bank_stats:
        return questions
    for q in questions:
        qid = str(q.get("id", ""))
        if not qid:
            continue
        stat = bank_stats.get(qid)
        if not stat:
            continue
        q["stats"] = {
            "total": int(stat.get("total", 0)),
            "correct": int(stat.get("correct", 0)),
            "rate": float(stat.get("rate", 0)),
        }
    return questions


def update_global_question_stats(bank_key: str, question_id: str, is_correct: bool):
    """更新全站题目统计"""
    if db_runtime_enabled():
        stat = db_storage.increment_question_stats(bank_key, question_id, is_correct)
        QUESTION_GLOBAL_STATS.setdefault(bank_key, {})[question_id] = stat
        update_cached_question_stats(bank_key, question_id, stat)
        return
    bank_stats = QUESTION_GLOBAL_STATS.setdefault(bank_key, {})
    current = bank_stats.setdefault(question_id, {"total": 0, "correct": 0, "rate": 0})
    current_total = int(current.get("total", 0)) + 1
    current_correct = int(current.get("correct", 0)) + (1 if is_correct else 0)
    current_rate = round(current_correct / current_total * 100, 1) if current_total else 0
    bank_stats[question_id] = {
        "total": current_total,
        "correct": current_correct,
        "rate": current_rate,
    }
    update_cached_question_stats(bank_key, question_id, bank_stats[question_id])


# ============== 工具函数 ==============

def normalize_judge_answer(value: Any) -> Optional[bool]:
    """将判断题答案统一转换为布尔值。无法识别时返回 None。"""
    if isinstance(value, bool):
        return value

    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if value == 1:
            return True
        if value == 0:
            return False

    if isinstance(value, str):
        text = value.strip().lower()
        if text in JUDGE_TRUE_VALUES:
            return True
        if text in JUDGE_FALSE_VALUES:
            return False

    return None


def normalize_blank_answer(value: Any) -> str:
    """将填空题答案统一为便于判分的文本。"""
    if value is None:
        return ""
    text = str(value)
    return re.sub(r"\s+", " ", text).strip()


def normalize_blank_answer_value(value: Any) -> Any:
    if isinstance(value, list):
        normalized = [
            normalize_blank_answer(item)
            for item in value
            if normalize_blank_answer(item)
        ]
        return normalized
    return normalize_blank_answer(value)


def is_blank_answer_correct(user_answer: Any, correct_answer: Any) -> bool:
    normalized_user = normalize_blank_answer(user_answer).casefold()
    if not normalized_user:
        return False

    if isinstance(correct_answer, list):
        candidates = correct_answer
    else:
        candidates = [correct_answer]

    return any(
        normalized_user == normalize_blank_answer(candidate).casefold()
        for candidate in candidates
    )


def infer_judge_answer_from_analysis(analysis: Any, allow_positive_fallback: bool = False) -> Optional[bool]:
    """从解析文本中推断判断题答案。"""
    if not isinstance(analysis, str):
        return None

    text = analysis.strip()
    if not text:
        return None

    head = text[:64]
    false_markers = [
        "错误", "不正确", "有误", "不准确", "不成立", "不符合", "并非",
        "不是", "不属于", "不能", "不存在", "不对",
    ]
    true_markers = ["正确", "无误", "成立", "准确", "符合", "确实", "属实"]

    if any(marker in head for marker in false_markers):
        return False
    if any(marker in head for marker in true_markers):
        return True

    if re.search(r"(因此|所以|由此可见|综上).{0,12}(错误|不正确|不准确)", text):
        return False
    if re.search(r"(因此|所以|由此可见|综上).{0,12}(正确|成立|准确)", text):
        return True

    if allow_positive_fallback and not any(marker in text for marker in false_markers):
        return True

    return None


def normalize_text_for_match(text: Any) -> str:
    if not isinstance(text, str):
        return ""
    return re.sub(r"[^0-9A-Za-z一-鿿]+", "", text).lower()


def extract_choice_letters(text: Any) -> Optional[List[str]]:
    if not isinstance(text, str):
        return None
    letters = re.findall(r"[A-F]", text.upper())
    if not letters:
        return None
    deduped: List[str] = []
    for letter in letters:
        if letter not in deduped:
            deduped.append(letter)
    return deduped or None


def choice_letters_to_answer(letters: List[str], q_type: str) -> Tuple[str, Any]:
    indices = [ord(letter) - 65 for letter in letters if 'A' <= letter <= 'F']
    if not indices:
        return q_type, None
    if q_type == "multi" or len(indices) > 1:
        return "multi", indices
    return "single", indices[0]


def infer_choice_answer_from_analysis(
    analysis: Any,
    options: Optional[List[str]],
    q_type: str,
    stem: str = "",
) -> Tuple[str, Any]:
    if not isinstance(analysis, str):
        return q_type, None

    text = analysis.strip()
    if not text:
        return q_type, None

    explicit_patterns = [
        r"正确答案(?:是|为|应为|应是)?[:：]?\s*([A-F](?:\s*[、,，和及]\s*[A-F])*)",
        r"答案(?:是|为|应为|应是)?[:：]?\s*([A-F](?:\s*[、,，和及]\s*[A-F])*)",
        r"([A-F](?:\s*[、,，和及]\s*[A-F])+)项均正确",
        r"([A-F](?:\s*[、,，和及]\s*[A-F])+)选项均正确",
        r"选项\s*([A-F](?:\s*[、,，和及]\s*[A-F])*)\s*(?:正确|符合|准确)",
    ]
    for pattern in explicit_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            letters = extract_choice_letters(match.group(1))
            if letters:
                return choice_letters_to_answer(letters, q_type)

    normalized_stem = normalize_text_for_match(stem)
    if any(marker in normalized_stem for marker in ["错误", "不正确", "不准确", "不属于", "不恰当"]):
        match = re.search(r"([A-F])项(?:错误|不正确|不准确|不符合|不属于)", text, re.I)
        if match:
            letters = extract_choice_letters(match.group(1))
            if letters:
                return choice_letters_to_answer(letters, "single")
    else:
        match = re.search(r"([A-F])项(?:正确|符合|准确)", text, re.I)
        if match:
            letters = extract_choice_letters(match.group(1))
            if letters:
                return choice_letters_to_answer(letters, q_type)

    if not options:
        return q_type, None

    normalized_analysis = normalize_text_for_match(text)
    if not normalized_analysis:
        return q_type, None

    negative_markers = [
        "其他选项", "其余选项", "均不", "错误", "不正确", "不准确",
        "不符合", "不选", "均非", "表述错误",
    ]
    negative_positions = [
        pos for pos in (normalized_analysis.find(normalize_text_for_match(marker)) for marker in negative_markers)
        if pos >= 0
    ]
    cutoff = min(negative_positions) if negative_positions else len(normalized_analysis)

    positions: List[Tuple[int, int]] = []
    for index, option in enumerate(options):
        normalized_option = normalize_text_for_match(option)
        if not normalized_option:
            continue
        pos = normalized_analysis.find(normalized_option)
        if pos >= 0:
            positions.append((index, pos))

    early_indices = [index for index, pos in positions if pos < cutoff and pos < 50]
    if early_indices:
        return choice_letters_to_answer([chr(index + 65) for index in early_indices], q_type)

    if len(positions) == 1:
        only_index = positions[0][0]
        return choice_letters_to_answer([chr(only_index + 65)], q_type)

    return q_type, None


def normalize_choice_answer(
    answer: Any,
    q_type: str,
    options: Optional[List[str]],
    analysis: Any,
    stem: str = "",
) -> Tuple[str, Any]:
    if isinstance(answer, list):
        normalized_indices = [int(item) for item in answer if isinstance(item, (int, float, str)) and str(item).isdigit()]
        if normalized_indices:
            if q_type == "multi" or len(normalized_indices) > 1:
                return "multi", normalized_indices
            return "single", normalized_indices[0]

    if isinstance(answer, (int, float)) and not isinstance(answer, bool):
        return ("multi", [int(answer)]) if q_type == "multi" else ("single", int(answer))

    letters = extract_choice_letters(answer)
    if letters:
        return choice_letters_to_answer(letters, q_type)

    return infer_choice_answer_from_analysis(analysis, options, q_type, stem)


def parse_question_bank(data: Dict, bank_key: str) -> List[Dict]:
    """解析题库为统一格式"""
    questions = []
    
    # 检测格式类型
    if "questions" in data:
        # 新格式：补齐 chapter/chapter_id 兼容字段
        chapter_to_id: Dict[str, str] = {}
        next_chapter_no = 1
        normalized: List[Dict] = []
        for i, raw_q in enumerate(data["questions"]):
            if not isinstance(raw_q, dict):
                continue
            q = dict(raw_q)
            q_type = _normalize_q_type(
                q.get("type", "single"),
                _answer_to_text(q.get("answer"), "single"),
            )
            q["type"] = q_type
            chapter_name = (
                q.get("chapter")
                or q.get("chapterName")
                or q.get("section")
                or q.get("group")
                or q.get("章节")
                or q.get("组别")
            )
            chapter_id = q.get("chapter_id")

            if chapter_name and not chapter_id:
                if chapter_name not in chapter_to_id:
                    chapter_to_id[chapter_name] = f"ch{next_chapter_no:02d}"
                    next_chapter_no += 1
                chapter_id = chapter_to_id[chapter_name]
                q["chapter_id"] = chapter_id
                q["chapter"] = chapter_name
            elif chapter_id and not chapter_name:
                q["chapter"] = str(chapter_id)

            # 保底 id
            if not q.get("id"):
                q["id"] = f"q{i+1:04d}"

            options = q.get("options")
            answer = q.get("answer")

            if isinstance(options, list) and len(options) > 0:
                q["options"] = options
            else:
                q["options"] = None

            answer_is_blank = answer is None or (isinstance(answer, str) and not answer.strip())

            if q_type == "blank":
                q["options"] = None
                q["answer"] = normalize_blank_answer_value(answer)
                if not q["answer"]:
                    continue
                normalized.append(q)
                continue

            if q_type in ("single", "multi") and q.get("options"):
                normalized_type, normalized_answer = normalize_choice_answer(
                    answer,
                    q_type,
                    q.get("options"),
                    q.get("analysis"),
                    q.get("content", ""),
                )
                q["type"] = normalized_type
                if normalized_answer is not None:
                    q["answer"] = normalized_answer

            # 仅兼容“被错误标成 single 的判断题”；multi 一律不转判断题
            if q_type == "single" and not q.get("options"):
                inferred_answer = normalize_judge_answer(answer)
                if inferred_answer is None and answer_is_blank:
                    inferred_answer = infer_judge_answer_from_analysis(
                        q.get("analysis"),
                        allow_positive_fallback=True,
                    )
                if inferred_answer is not None:
                    q["type"] = "judge"
                    q["answer"] = inferred_answer
                    q["options"] = None

            if q.get("type") == "judge":
                normalized_answer = normalize_judge_answer(q.get("answer"))
                if normalized_answer is None and answer_is_blank:
                    normalized_answer = infer_judge_answer_from_analysis(
                        q.get("analysis"),
                        allow_positive_fallback=True,
                    )
                if normalized_answer is not None:
                    q["answer"] = normalized_answer
                q["options"] = None

            # 题目数据异常：选择题缺失选项或答案未能规范化时，直接跳过
            if q.get("type") in ("single", "multi") and not q.get("options"):
                continue
            if q.get("type") == "single" and not isinstance(q.get("answer"), int):
                continue
            if q.get("type") == "multi" and not isinstance(q.get("answer"), list):
                continue

            normalized.append(q)
        return apply_global_question_stats(bank_key, normalized)
    
    # 旧格式转换
    qid = 1
    chapter_no = 1
    for chapter_name, types in data.items():
        chapter_id = f"ch{chapter_no:02d}"
        chapter_no += 1
        
        for type_name, items in types.items():
            q_type = "single"
            if "多选" in type_name or "多项" in type_name:
                q_type = "multi"
            elif "判断" in type_name:
                q_type = "judge"
            elif "填空" in type_name or "填充" in type_name:
                q_type = "blank"
            
            for item_text in items:
                question = parse_question_text(item_text, q_type, chapter_id, chapter_name, qid)
                if question:
                    questions.append(question)
                    qid += 1
    
    return apply_global_question_stats(bank_key, questions)


def parse_question_text(text: str, q_type: str, chapter_id: str, chapter_name: str, qid: int) -> Optional[Dict]:
    """解析单个题目文本"""
    lines = text.strip().split('\n')
    if not lines:
        return None
    
    # 提取题干
    stem = lines[0]
    # 去掉题号
    stem = re.sub(r'^\d+[、.]\s*', '', stem)
    
    # 提取选项
    options = []
    answer = None
    analysis = ""
    
    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        
        # 答案行
        answer_match = re.match(r'^答案\s*[：:]\s*(.+)$', line)
        if answer_match:
            answer_text = answer_match.group(1).strip()
            
            if q_type == "blank":
                answer = normalize_blank_answer(answer_text)
            elif q_type == "judge":
                answer = answer_text.lower() in ["对", "正确", "√", "true", "t", "是", "yes", "y"]
            elif q_type == "multi":
                # 多选答案转为索引数组
                answer = [ord(c.upper()) - 65 for c in answer_text if c.upper() in 'ABCDEF']
            else:
                # 单选转为索引
                m = re.search(r'[A-Fa-f]', answer_text)
                answer = ord(m.group(0).upper()) - 65 if m else 0
        
        # 解析行
        elif re.match(r'^解析\s*[：:]', line):
            analysis = re.sub(r'^解析\s*[：:]\s*', '', line).strip()
        
        # 选项行
        elif re.match(r'^[A-F][、.．]', line):
            option_text = re.sub(r'^[A-F][、.．]\s*', '', line)
            options.append(option_text)
    
    return {
        "id": f"q{qid:04d}",
        "type": q_type,
        "chapter": chapter_name,
        "chapter_id": chapter_id,
        "content": stem,
        "options": options if options else None,
        "answer": answer,
        "analysis": analysis,
        "stats": {
            "total": 0,
            "correct": 0,
            "rate": 0
        }
    }


# ============== FastAPI 应用 ==============

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    print("🚀 正在初始化...")
    initialize_database_if_configured()
    load_question_banks()
    sync_question_banks_to_db()
    load_rankings()
    load_question_stats()
    load_food_wheel_items()
    refresh_question_cache()
    print(f"📚 已加载 {len(QUESTION_BANKS)} 个题库")
    yield
    print("👋 正在关闭...")
    save_rankings()
    save_question_stats()


app = FastAPI(
    title="刷题系统 API",
    description="React + TypeScript 刷题应用后端",
    version="2.0.0",
    lifespan=lifespan
)

# CORS 配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=get_cors_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============== API 路由 ==============

@app.get("/api/admin/session")
async def get_admin_session_status(request: Request):
    """Return browser admin-session state without exposing the session cookie."""
    authenticated = is_admin_session_valid(
        request.cookies.get(ADMIN_SESSION_COOKIE_NAME)
    )
    return {"authenticated": authenticated}


@app.post("/api/admin/session")
async def create_admin_session(
    request: Request,
    response: Response,
    x_admin_token: Optional[str] = Header(None),
):
    """Exchange the server-side admin secret for a short-lived HttpOnly session."""
    if not get_admin_token():
        raise HTTPException(status_code=503, detail="后台管理 Token 未配置")
    if not is_admin_token_valid(x_admin_token):
        raise HTTPException(status_code=403, detail="后台管理 Token 无效")

    ttl = get_admin_session_ttl_seconds()
    response.set_cookie(
        key=ADMIN_SESSION_COOKIE_NAME,
        value=create_admin_session_token(),
        max_age=ttl,
        httponly=True,
        secure=should_use_secure_admin_cookie(request),
        samesite="strict",
        path="/",
    )
    return {"authenticated": True, "expires_in": ttl}


@app.delete("/api/admin/session")
async def delete_admin_session(request: Request, response: Response):
    response.delete_cookie(
        key=ADMIN_SESSION_COOKIE_NAME,
        httponly=True,
        secure=should_use_secure_admin_cookie(request),
        samesite="strict",
        path="/",
    )
    return {"authenticated": False}

@app.get("/api/banks")
async def get_banks():
    """获取题库列表"""
    return {
        "banks": [
            build_bank_summary(key, bank)
            for key, bank in QUESTION_BANKS.items()
            if is_bank_enabled(key)
        ]
    }


@app.post("/api/practice/start")
async def start_practice(request: StartPracticeRequest):
    """开始练习"""
    require_enabled_bank(request.bank)

    questions = get_bank_questions(request.bank)
    
    mode = request.mode
    params = request.params
    raw_count = params.get("count", 20)
    try:
        count = int(raw_count)
    except (TypeError, ValueError):
        count = 20
    unlimited = count <= 0
    
    import random

    def _sample(pool: List[Dict]) -> List[Dict]:
        if not pool:
            return []
        if unlimited:
            return random.sample(pool, len(pool))
        return random.sample(pool, min(count, len(pool)))
    
    if mode == "random":
        # 随机模式
        selected = _sample(questions)
    
    elif mode == "chapter":
        # 章节模式
        chapter_id = params.get("chapter_id")
        chapter_q = [q for q in questions if q.get("chapter_id") == chapter_id or q.get("chapter") == chapter_id]
        selected = random.sample(chapter_q, len(chapter_q)) if chapter_q else []
    
    elif mode == "hard":
        # 难题模式
        threshold = params.get("threshold", 50)
        hard_q = [q for q in questions if q.get("stats", {}).get("total", 0) > 0 and q.get("stats", {}).get("rate", 100) < threshold]
        if not hard_q:
            # 如果没有统计，随机选
            hard_q = questions
        hard_q.sort(key=lambda x: x.get("stats", {}).get("rate", 100))
        selected = hard_q if unlimited else hard_q[:count]
    
    else:
        selected = _sample(questions)
    
    # 计算平均正确率
    avg_rate = sum(q.get("stats", {}).get("rate", 0) for q in selected) / len(selected) if selected else 0
    
    return {
        "questions": selected,
        "total": len(selected),
        "avg_rate": round(avg_rate, 1)
    }


@app.post("/api/practice/submit")
async def submit_answer(request: SubmitAnswerRequest):
    """提交答案"""
    require_enabled_bank(request.bank)

    question = get_bank_question(request.bank, request.question_id)
    if not question:
        raise HTTPException(status_code=404, detail="题目不存在")
    
    # 检查答案
    user_answer = request.answer
    correct_answer = question["answer"]
    q_type = question["type"]
    
    is_correct = False
    response_correct_answer = correct_answer
    if q_type == "judge":
        normalized_user_answer = normalize_judge_answer(user_answer)
        normalized_correct_answer = normalize_judge_answer(correct_answer)
        if normalized_correct_answer is not None:
            response_correct_answer = normalized_correct_answer
        is_correct = (
            normalized_user_answer is not None
            and normalized_correct_answer is not None
            and normalized_user_answer == normalized_correct_answer
        )
    elif q_type == "blank":
        is_correct = is_blank_answer_correct(user_answer, correct_answer)
    elif q_type == "multi":
        user_sorted = sorted(user_answer) if isinstance(user_answer, list) else []
        correct_sorted = sorted(correct_answer) if isinstance(correct_answer, list) else []
        is_correct = user_sorted == correct_sorted
    else:
        is_correct = user_answer == correct_answer
    
    # 更新用户统计
    resolved_user_id = request.user_id
    user_stats_payload: Optional[Dict[str, Any]] = None
    if request.user_id:
        current_name = (USER_STATS.get(request.user_id, {}).get("name") or request.user_id)
        if db_runtime_enabled():
            updated_stats = db_storage.increment_user_stats(
                request.user_id,
                current_name,
                is_correct,
            )
            resolved_user_id = str(updated_stats.get("user_id") or request.user_id)
            resolved_name = str(updated_stats.get("name") or current_name)
            if resolved_user_id != request.user_id:
                if request.user_id in USER_STATS:
                    USER_STATS.pop(request.user_id, None)
                NAME_TO_ID[resolved_name] = resolved_user_id
            USER_STATS.setdefault(resolved_user_id, {
                "name": resolved_name,
                "correct": 0,
                "total": 0,
                "practice_history": [],
            }).update({
                "name": resolved_name,
                "correct": updated_stats["correct"],
                "total": updated_stats["total"],
            })
            user_stats_payload = {
                "user_id": resolved_user_id,
                "name": resolved_name,
                "correct": updated_stats["correct"],
                "total": updated_stats["total"],
                "rate": round(
                    updated_stats["correct"] / max(1, updated_stats["total"]) * 100,
                    1,
                ),
            }
        else:
            USER_STATS.setdefault(request.user_id, {
                "name": current_name,
                "correct": 0,
                "total": 0,
                "practice_history": [],
            })
            USER_STATS[request.user_id]["total"] += 1
            if is_correct:
                USER_STATS[request.user_id]["correct"] += 1
            resolved_user_id = request.user_id
            save_rankings()
            resolved_user_stats = USER_STATS[resolved_user_id]
            user_stats_payload = {
                "user_id": resolved_user_id,
                "name": resolved_user_stats.get("name", resolved_user_id),
                "correct": resolved_user_stats["correct"],
                "total": resolved_user_stats["total"],
                "rate": round(
                    resolved_user_stats["correct"]
                    / max(1, resolved_user_stats["total"]) * 100,
                    1,
                ),
            }

    # 更新全站题目统计
    update_global_question_stats(request.bank, request.question_id, is_correct)
    save_question_stats()
    
    return {
        "correct": is_correct,
        "correct_answer": response_correct_answer,
        "analysis": question.get("analysis", ""),
        "stats": question.get("stats", {}),
        "user_stats": {
            "user_id": user_stats_payload["user_id"],
            "name": user_stats_payload["name"],
            "correct": user_stats_payload["correct"],
            "total": user_stats_payload["total"],
            "rate": user_stats_payload["rate"],
        } if user_stats_payload is not None else None
    }


@app.post("/api/user")
async def set_user(request: UserRequest):
    """设置用户信息"""
    global NEXT_USER_ID
    
    name = request.name or ""
    
    # 检查是否已存在
    if name and name in NAME_TO_ID:
        user_id = NAME_TO_ID[name]
        return {
            "user_id": user_id,
            "name": name,
            **USER_STATS[user_id]
        }

    if db_runtime_enabled() and name:
        found = db_storage.find_user_by_name(name)
        if found:
            user_id, stats = found
            USER_STATS[user_id].update(stats)
            NAME_TO_ID[name] = user_id
            return {
                "user_id": user_id,
                "name": name,
                **USER_STATS[user_id]
            }
    
    # 创建新用户。系统生成 ID 使用前缀，避免与学号/工号这类纯数字自定义 ID 混淆。
    user_id = allocate_generated_user_id()
    
    USER_STATS[user_id] = {
        "name": name or user_id,
        "correct": 0,
        "total": 0,
        "practice_history": []
    }
    
    if name:
        NAME_TO_ID[name] = user_id
    
    if db_runtime_enabled():
        db_stats = db_storage.upsert_user(user_id, name or user_id)
        USER_STATS[user_id].update(db_stats)
    else:
        save_rankings()
    
    return {
        "user_id": user_id,
        "name": name or user_id,
        "correct": 0,
        "total": 0
    }


@app.get("/api/ranking")
async def get_ranking():
    """获取排行榜"""
    if db_runtime_enabled():
        return {"ranking": db_storage.get_ranking()}

    ranking = []
    for user_id, stats in USER_STATS.items():
        if stats["total"] > 0:
            ranking.append({
                "user_id": user_id,
                "name": stats.get("name", user_id),
                "correct": stats["correct"],
                "total": stats["total"],
                "accuracy": round(stats["correct"] / stats["total"] * 100, 1)
            })

    ranking.sort(key=lambda x: (-x["correct"], -x["accuracy"]))
    return {"ranking": ranking}


@app.get("/api/wheel")
async def get_food_wheel():
    wheels = load_food_wheel_items()
    if not wheels:
        wheels = []
    return {
        "wheels": wheels,
    }


@app.post("/api/wheel")
async def update_food_wheel(request: FoodWheelRequest):
    owner_user_id = (request.user_id or "").strip()
    if not owner_user_id:
        raise HTTPException(
            status_code=422,
            detail="上传美食转盘需要 user_id",
        )

    normalized = _normalize_food_wheel_items(request.items)
    if len(normalized) < 2:
        raise HTTPException(
            status_code=422,
            detail="美食转盘至少需要 2 个选项",
        )

    saved = save_food_wheel_items(owner_user_id, normalized)
    return {
        **saved,
    }


@app.post("/api/feedback")
async def create_feedback(request: FeedbackRequest):
    question_index = request.question_index
    suggestion = _normalize_feedback_suggestion(request.suggestion)
    question_id = (request.question_id or "").strip() or None
    question_content = (request.question_content or "").strip() or None

    if question_index <= 0:
        raise HTTPException(status_code=422, detail="题目索引必须是大于 0 的整数")
    if not suggestion:
        raise HTTPException(status_code=422, detail="建议改正内容不能为空")
    if len(suggestion) > 2000:
        raise HTTPException(status_code=422, detail="建议改正内容不能超过 2000 字符")
    if question_content and len(question_content) > 2000:
        question_content = question_content[:2000]

    source_page = (request.source_page or "quiz").strip() or "quiz"
    question_bank = _normalize_feedback_bank(request.question_bank)

    if db_runtime_enabled():
        try:
            result = db_storage.create_feedback(
                question_index=question_index,
                suggestion=suggestion,
                question_bank=question_bank,
                question_id=question_id,
                question_content=question_content,
                source_page=source_page,
            )
            return {
                "ok": True,
                "feedback_id": result["feedback_id"],
                "question_index": result["question_index"],
                "question_id": result.get("question_id"),
                "question_bank": result.get("question_bank"),
                "created_at": result["created_at"],
            }
        except Exception as e:
            print(f"保存反馈到 PostgreSQL 失败: {e}")

    result = save_feedback_fallback(
        question_index=question_index,
        suggestion=suggestion,
        question_bank=question_bank,
        question_id=question_id,
        question_content=question_content,
        source_page=source_page,
    )
    return {
        "ok": True,
        "feedback_id": result["feedback_id"],
        "question_index": result["question_index"],
        "question_id": result.get("question_id"),
        "question_bank": result.get("question_bank"),
        "created_at": result["created_at"],
    }


@app.get("/api/feedback/dashboard")
async def get_feedback_dashboard():
    pending_limit = 100
    resolved_limit = 100
    archived_limit = 100
    if db_runtime_enabled():
        try:
            today_start, today_end = _feedback_today_range()
            return db_storage.get_feedback_dashboard(
                today_start=today_start,
                today_end=today_end,
                pending_limit=pending_limit,
                resolved_limit=resolved_limit,
                archived_limit=archived_limit,
            )
        except Exception as e:
            print(f"加载反馈看板失败: {e}")

    return load_feedback_dashboard_fallback(
        pending_limit=pending_limit,
        resolved_limit=resolved_limit,
        archived_limit=archived_limit,
    )


@app.patch("/api/feedback/{feedback_id}/status", dependencies=[Depends(require_admin_token)])
async def update_feedback_status(feedback_id: int, request: FeedbackStatusRequest):
    status = _normalize_feedback_status(request.status)
    note = (request.resolution_note or "").strip()
    if len(note) > 1000:
        raise HTTPException(status_code=422, detail="处理备注不能超过 1000 字符")

    item = None
    if db_runtime_enabled():
        try:
            item = db_storage.update_feedback_status(
                feedback_id=feedback_id,
                status=status,
                resolution_note=note,
            )
        except Exception as e:
            print(f"更新反馈处理状态失败: {e}")
            raise HTTPException(status_code=500, detail="更新反馈处理状态失败")
    else:
        item = update_feedback_status_fallback(
            feedback_id=feedback_id,
            status=status,
            resolution_note=note,
        )

    if not item:
        raise HTTPException(status_code=404, detail="反馈不存在")
    return {
        "ok": True,
        "item": item,
    }


# ============== 文件提取 API ==============

@app.post("/api/extract/parse", dependencies=[Depends(require_admin_token)])
async def extract_parse(file: UploadFile = File(...)):
    """解析上传的文件"""
    # 保存临时文件
    suffix = os.path.splitext(file.filename)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
    
    try:
        # 根据文件类型提取文本
        text = ""
        suffix_lower = suffix.lower()
        if suffix_lower == '.pdf':
            text = extract_text_from_pdf(tmp_path)
            questions = parse_questions_from_text(text)
        elif suffix_lower in ['.docx', '.doc']:
            text = extract_text_from_docx(tmp_path)
            questions = parse_questions_from_docx(tmp_path)
        elif suffix_lower == '.json':
            text = content.decode('utf-8-sig')
            questions = parse_questions_from_json_text(text)
            if not questions:
                raise HTTPException(status_code=400, detail="JSON 文件中未识别到题目，请检查字段或结构")
            return {
                "content": text[:1000] + "..." if len(text) > 1000 else text,
                "questions": questions
            }
        else:
            text = content.decode('utf-8-sig')
            questions = parse_questions_from_text(text)

        return {
            "content": text[:1000] + "..." if len(text) > 1000 else text,
            "questions": questions
        }
    finally:
        os.unlink(tmp_path)


def extract_text_from_pdf(path: str) -> str:
    """从 PDF 提取文本"""
    try:
        import PyPDF2
        text = ""
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        raise HTTPException(status_code=500, detail="需要安装 PyPDF2: pip install PyPDF2")


def extract_text_from_docx(path: str) -> str:
    """从 Word 提取文本"""
    try:
        from docx import Document
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except ImportError:
        raise HTTPException(status_code=500, detail="需要安装 python-docx: pip install python-docx")


def parse_questions_from_docx(path: str) -> List[Dict]:
    """按段落结构解析 Word 题库，避免题干中的 IP/版本号被误判成新题号。"""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="需要安装 python-docx: pip install python-docx")

    question_start_pattern = re.compile(r'^(?P<number>\d+)[、.．]\s*(?P<content>.+)$')
    option_pattern = re.compile(r'^(?P<label>[A-F])[、.．]\s*(?P<text>.+)$')
    answer_pattern = re.compile(r'^答案\s*[：:]\s*(?P<answer>.+)$')
    analysis_pattern = re.compile(r'^(?:解析|答案解析)\s*[：:]\s*(?P<analysis>.*)$')
    type_pattern = re.compile(r'【(?P<type>单选题|多选题|判断题|填空题)】')

    doc = Document(path)
    paragraphs = [re.sub(r'\s+', ' ', p.text).strip() for p in doc.paragraphs if p.text and p.text.strip()]
    questions: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None
    analysis_lines: List[str] = []
    chapter_name = "默认章节"
    chapter_id = "ch01"

    def finalize_current():
        nonlocal current, analysis_lines
        if not current:
            return
        current["content"] = re.sub(r'\s+', ' ', current["content"]).strip()
        current["options"] = [re.sub(r'\s+', ' ', option).strip() for option in current["options"] if option.strip()]
        current["analysis"] = re.sub(r'\s+', ' ', " ".join(analysis_lines or [current["analysis"]])).strip()
        if current["type"] == "judge" and not current["options"]:
            current["options"] = None
        questions.append(current)
        current = None
        analysis_lines = []

    def ensure_current(number: str, raw_content: str):
        nonlocal current, analysis_lines
        finalize_current()
        content = raw_content.strip()
        q_type = "single"
        type_match = type_pattern.search(content)
        if type_match:
            type_map = {"单选题": "single", "多选题": "multi", "判断题": "judge", "填空题": "blank"}
            q_type = type_map[type_match.group("type")]
            content = type_pattern.sub("", content, count=1).strip()
        current = {
            "id": f"q{len(questions) + 1:04d}",
            "number": str(number),
            "type": q_type,
            "content": content,
            "options": [],
            "answer": "",
            "analysis": "",
            "chapter": chapter_name,
            "chapter_id": chapter_id,
        }
        analysis_lines = []

    for paragraph in paragraphs:
        question_match = question_start_pattern.match(paragraph)
        if question_match:
            ensure_current(question_match.group("number"), question_match.group("content"))
            continue

        if current is None:
            continue

        answer_match = answer_pattern.match(paragraph)
        if answer_match:
            answer_text = answer_match.group("answer").strip()
            answer_letters = extract_choice_letters(answer_text)
            if current["type"] in {"single", "multi"} and answer_letters:
                current["answer"] = "".join(answer_letters)
                if len(answer_letters) > 1:
                    current["type"] = "multi"
            else:
                current["answer"] = answer_text
            continue

        analysis_match = analysis_pattern.match(paragraph)
        if analysis_match:
            analysis_head = analysis_match.group("analysis").strip()
            analysis_lines = [analysis_head] if analysis_head else []
            continue

        option_match = option_pattern.match(paragraph)
        if option_match:
            current["options"].append(option_match.group("text").strip())
            continue

        if analysis_lines:
            analysis_lines.append(paragraph)
            continue

        if current["options"]:
            current["options"][-1] = f"{current['options'][-1]} {paragraph}".strip()
            continue

        current["content"] = f"{current['content']} {paragraph}".strip()

    finalize_current()
    return questions


def _answer_to_text(answer: Any, q_type: str) -> str:
    """将不同题库答案格式统一为文本格式。"""
    if answer is None:
        return ""
    if isinstance(answer, bool):
        return "对" if answer else "错"
    if isinstance(answer, int):
        if q_type in {"single", "multi"} and 0 <= answer < 26:
            return chr(65 + answer)
        return str(answer)
    if isinstance(answer, list):
        letters = []
        for item in answer:
            if isinstance(item, int) and 0 <= item < 26:
                letters.append(chr(65 + item))
            else:
                letters.append(str(item))
        return "".join(letters)
    return str(answer)


def _normalize_q_type(raw_type: Any, answer_text: str = "") -> str:
    type_text = str(raw_type or "").strip().lower()
    if type_text in {"single", "radio", "choice", "single_choice"}:
        return "single"
    if type_text in {"multi", "multiple", "multiple_choice", "checkbox"}:
        return "multi"
    if type_text in {"judge", "tf", "truefalse", "boolean"}:
        return "judge"
    if type_text in {"blank", "fill", "fill_blank", "fill_in", "completion"}:
        return "blank"

    # 中文题型兼容
    if any(k in str(raw_type) for k in ["多选", "多项"]):
        return "multi"
    if any(k in str(raw_type) for k in ["判断", "是非"]):
        return "judge"
    if any(k in str(raw_type) for k in ["填空", "填充", "补全"]):
        return "blank"
    if any(k in str(raw_type) for k in ["单选", "单项"]):
        return "single"

    # 由答案兜底推断
    if answer_text and len(answer_text) > 1 and all(c.upper() in "ABCDEF" for c in answer_text):
        return "multi"
    if answer_text in {"对", "错", "正确", "错误", "√", "×", "true", "false", "True", "False"}:
        return "judge"
    return "single"


def _normalize_options(raw_options: Any) -> List[str]:
    if raw_options is None:
        return []
    if isinstance(raw_options, list):
        return [str(x).strip() for x in raw_options]
    if isinstance(raw_options, dict):
        # 兼容 {"A":"xxx","B":"yyy"}
        items = sorted(raw_options.items(), key=lambda kv: str(kv[0]))
        return [str(v).strip() for _, v in items]
    return []


def _normalize_question_item(raw: Dict, qid: int) -> Dict:
    content = (
        raw.get("content")
        or raw.get("question")
        or raw.get("title")
        or raw.get("stem")
        or raw.get("text")
        or raw.get("题目")
        or raw.get("题干")
        or ""
    )
    raw_answer = raw.get(
        "answer",
        raw.get("correctAnswer", raw.get("correct_answer", raw.get("答案", "")))
    )
    prelim_answer = _answer_to_text(raw_answer, "single")
    q_type = _normalize_q_type(
        raw.get("type", raw.get("questionType", raw.get("q_type", raw.get("题型", "")))),
        prelim_answer
    )
    answer_text = _answer_to_text(raw_answer, q_type)
    options = _normalize_options(
        raw.get(
            "options",
            raw.get("choices", raw.get("option", raw.get("optionList", raw.get("选项"))))
        )
    )
    return {
        "id": str(raw.get("id", f"q{qid:04d}")),
        "number": str(raw.get("number", qid)),
        "type": q_type,
        "content": str(content).strip(),
        "options": options or [],
        "answer": answer_text,
        "analysis": str(raw.get("analysis", raw.get("explanation", raw.get("解析", raw.get("答案解析", ""))))).strip(),
        "chapter": raw.get(
            "chapter",
            raw.get(
                "chapterName",
                raw.get(
                    "section",
                    raw.get("group", raw.get("章节", raw.get("组别"))),
                ),
            ),
        ),
    }


def parse_questions_from_json_text(text: str) -> List[Dict]:
    """从 JSON 文本解析题目，兼容多种题库结构。"""
    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"JSON 格式错误: {e.msg}")

    def _extract_from_list(items: Any) -> List[Dict]:
        if not isinstance(items, list):
            return []
        return [
            _normalize_question_item(item, idx + 1)
            for idx, item in enumerate(items)
            if isinstance(item, dict)
        ]

    # 1) 标准结构: {"questions": [...]}
    if isinstance(data, dict) and isinstance(data.get("questions"), list):
        return _extract_from_list(data["questions"])

    # 1.1) 常见变体: {"data": {"questions": [...]}} / {"items":[...]} / {"list":[...]}
    if isinstance(data, dict):
        for key in ["items", "list", "records", "problemList", "题目", "试题", "questionList", "dataList"]:
            if isinstance(data.get(key), list):
                return _extract_from_list(data[key])
        nested = data.get("data")
        if isinstance(nested, dict):
            for key in ["questions", "items", "list", "records", "problemList", "题目", "试题", "questionList", "dataList"]:
                if isinstance(nested.get(key), list):
                    return _extract_from_list(nested[key])
        if isinstance(nested, list):
            return _extract_from_list(nested)

    # 1.2) 章节数组: {"chapters":[{"name":"xx","questions":[...]}]}
    if isinstance(data, dict) and isinstance(data.get("chapters"), list):
        questions: List[Dict] = []
        qid = 1
        for chapter in data["chapters"]:
            if not isinstance(chapter, dict):
                continue
            chapter_name = chapter.get("name", chapter.get("chapter", ""))
            chapter_questions = chapter.get("questions", chapter.get("items", []))
            if not isinstance(chapter_questions, list):
                continue
            for item in chapter_questions:
                if not isinstance(item, dict):
                    continue
                norm = _normalize_question_item(item, qid)
                if chapter_name and not norm.get("chapter"):
                    norm["chapter"] = chapter_name
                questions.append(norm)
                qid += 1
        if questions:
            return questions

    # 2) 直接数组: [...]
    if isinstance(data, list):
        return _extract_from_list(data)

    # 3) 旧章节结构: {chapter: {type: [text, ...]}}
    def _is_legacy_chapter_shape(obj: Any) -> bool:
        if not isinstance(obj, dict) or not obj:
            return False
        for v in obj.values():
            if not isinstance(v, dict) or not v:
                return False
            if not all(isinstance(items, list) for items in v.values()):
                return False
        return True

    if isinstance(data, dict) and _is_legacy_chapter_shape(data):
        legacy_questions = parse_question_bank(data, "uploaded_json")
        return [
            _normalize_question_item(item, idx + 1)
            for idx, item in enumerate(legacy_questions)
            if isinstance(item, dict)
        ]

    # 4) 最后兜底：在字典中递归查找“看起来像题目列表”的数组
    def _looks_like_question(item: Any) -> bool:
        if not isinstance(item, dict):
            return False
        keys = set(item.keys())
        candidates = {
            "content", "question", "title", "stem", "text", "题目", "题干",
            "options", "choices", "optionList", "选项",
            "answer", "correctAnswer", "correct_answer", "答案"
        }
        return len(keys & candidates) >= 2

    def _find_question_list(node: Any) -> Optional[List[Dict]]:
        if isinstance(node, list) and node and all(isinstance(x, dict) for x in node):
            if sum(1 for x in node if _looks_like_question(x)) >= max(1, len(node) // 3):
                return _extract_from_list(node)
        if isinstance(node, dict):
            for v in node.values():
                found = _find_question_list(v)
                if found:
                    return found
        return None

    if isinstance(data, dict):
        found = _find_question_list(data)
        if found:
            return found

    raise HTTPException(status_code=400, detail="不支持的 JSON 题库结构，请提供 questions/items/list 或章节结构")


def _strip_source_noise(text: str) -> str:
    cleaned = text.replace('\r', '\n')
    cleaned = re.sub(r'---\s*PAGE\s+\d+\s*---', '\n', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'更多(?:考试)?资料请加[^\n]*', '\n', cleaned)
    cleaned = re.sub(r'更多考试资料请加[^\n]*', '\n', cleaned)
    cleaned = re.sub(r'河南大学考试墙[^\n]*', '\n', cleaned)
    cleaned = re.sub(r'河南大学小过儿[^\n]*', '\n', cleaned)
    cleaned = re.sub(r'严禁任何个人组织商家盗用或售卖！?', '\n', cleaned)
    return cleaned


def _is_source_noise_line(line: str) -> bool:
    value = re.sub(r'\s+', '', line or "")
    if not value:
        return True
    if re.fullmatch(r'\d{1,3}', value):
        return True
    if re.fullmatch(r'[~*\-—_]+', value):
        return True
    if not re.search(r'[\u4e00-\u9fff]', value) and len(value) <= 24:
        return True
    return False


def _clean_source_lines(text: str) -> str:
    return "\n".join(
        line.strip()
        for line in _strip_source_noise(text).splitlines()
        if not _is_source_noise_line(line)
    )


def _normalize_structured_chapter(raw: str) -> str:
    value = re.sub(r'\s+', '', raw or "")
    if value.startswith("绪论") or value.startswith("导论"):
        return "导论"
    if value.startswith("第一章"):
        return "第一章"
    if value.startswith("第二章"):
        return "第二章"
    if value.startswith("第三章"):
        return "第三章"
    if value.startswith("第四章"):
        return "第四章"
    if value.startswith("第五章"):
        return "第五章"
    if value.startswith("第六章") or value.startswith("第七章"):
        return "第六章第七章"
    if value.startswith("所有章节"):
        return "所有章节判断题"
    return raw.strip() or "默认章节"


def _parse_answer_key(answer_area: str) -> Tuple[Dict[str, Dict[str, Any]], List[bool]]:
    chapters = ["绪论", "导论", "第一章", "第二章", "第三章", "第四章", "第五章", "第六章", "第七章", "第六章第七章"]
    parts = re.split("(" + "|".join(chapters) + ")", answer_area.replace("：", ":"))
    answer_data: Dict[str, Dict[str, Any]] = {}
    current_chapter: Optional[str] = None

    for part in parts:
        value = part.strip()
        if not value:
            continue
        if value in chapters:
            current_chapter = _normalize_structured_chapter(value)
            answer_data.setdefault(current_chapter, {"single": "", "multi": []})
            continue
        if not current_chapter:
            continue

        single_match = re.search(r'单选\s*:+\s*', value)
        if single_match:
            after_single = value[single_match.end():]
            stop_positions = [
                match.start()
                for match in [
                    re.search(r'多选\s*:+', after_single),
                    re.search(r'判断\s+', after_single),
                ]
                if match
            ]
            single_raw = after_single[: min(stop_positions) if stop_positions else len(after_single)]
            answer_data[current_chapter]["single"] = "".join(re.findall(r'[A-D]', single_raw.upper()))

        multi_match = re.search(r'多选\s*:+\s*', value)
        if multi_match:
            after_multi = value[multi_match.end():]
            judge_match = re.search(r'判断\s+', after_multi)
            multi_raw = after_multi[: judge_match.start() if judge_match else len(after_multi)]
            multi_raw = re.sub(r'\d+\s*-\s*\d+', ' ', multi_raw)
            answer_data[current_chapter]["multi"] = re.findall(r'[A-D]+', multi_raw.upper())

    judge_answers: List[bool] = []
    judge_match = re.search(r'判断\s+', answer_area)
    if judge_match:
        judge_answers = [item == "对" for item in re.findall(r'[对错]', answer_area[judge_match.end():])]

    return answer_data, judge_answers


def _split_structured_sections(question_area: str) -> List[Tuple[str, str]]:
    chapter_pattern = re.compile(
        r'(?m)^\s*(导论|绪论|第[一二三四五六七]章[^\n]*|所有章节判断题汇总)\s*$'
    )
    matches = list(chapter_pattern.finditer(question_area))
    sections: List[Tuple[str, str]] = []
    for index, match in enumerate(matches):
        chapter = _normalize_structured_chapter(match.group(1))
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(question_area)
        sections.append((chapter, question_area[start:end]))
    return sections


def _split_structured_choice_parts(section_text: str) -> Tuple[str, str]:
    single_match = re.search(r'一\s*[、.．]\s*单项?选择题', section_text)
    multi_match = re.search(r'二\s*[、.．]\s*多项?选择题', section_text)
    if not single_match and not multi_match:
        return "", ""

    single_text = ""
    multi_text = ""
    if single_match:
        single_start = single_match.end()
        single_end = multi_match.start() if multi_match else len(section_text)
        single_text = section_text[single_start:single_end]
    if multi_match:
        multi_text = section_text[multi_match.end():]
    return single_text, multi_text


def _extract_structured_options(body: str) -> Tuple[str, List[str]]:
    cleaned = _clean_source_lines(body)
    option_pattern = re.compile(
        r'(?<![A-Za-z])([A-D])(?:[.．、]|\s+|(?=[\u4e00-\u9fff“《]))'
    )
    matches = list(option_pattern.finditer(cleaned))
    if len(matches) < 2:
        return re.sub(r'\s+', ' ', cleaned).strip(), []

    first_option = matches[0].start()
    stem = re.sub(r'\s+', ' ', cleaned[:first_option]).strip()
    options: List[str] = []

    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(cleaned)
        option = cleaned[start:end].strip()
        option = "\n".join(
            line.strip()
            for line in option.splitlines()
            if not _is_source_noise_line(line)
        )
        option = re.sub(r'\s+', ' ', option).strip()
        if option:
            options.append(option)

    return stem, options


def _parse_structured_choice_questions(section_text: str, chapter: str, q_type: str, qid_start: int) -> Tuple[List[Dict], int]:
    questions: List[Dict[str, Any]] = []
    starts = list(re.finditer(r'(?m)^\s*(\d+)[.．、）\)]\s*', section_text))
    for index, match in enumerate(starts):
        number = match.group(1)
        start = match.end()
        end = starts[index + 1].start() if index + 1 < len(starts) else len(section_text)
        body = section_text[start:end]
        stem, options = _extract_structured_options(body)
        if not stem or len(options) < 2:
            continue
        questions.append({
            "id": f"q{qid_start + len(questions):04d}",
            "number": number,
            "type": q_type,
            "content": stem,
            "options": options[:4],
            "answer": "",
            "analysis": "",
            "chapter": chapter,
            "chapter_id": f"ch_{chapter}",
            "stats": {"total": 0, "correct": 0, "rate": 0},
        })
    return questions, qid_start + len(questions)


def _parse_structured_judge_questions(section_text: str, qid_start: int) -> Tuple[List[Dict], int]:
    questions: List[Dict[str, Any]] = []
    cleaned = _clean_source_lines(section_text)
    starts = list(re.finditer(r'(?m)^\s*(\d+)[.．、）\)]{1,2}\s*', cleaned))
    for index, match in enumerate(starts):
        number = match.group(1)
        start = match.end()
        end = starts[index + 1].start() if index + 1 < len(starts) else len(cleaned)
        content = re.sub(r'\s+', ' ', cleaned[start:end]).strip()
        if not content:
            continue
        questions.append({
            "id": f"q{qid_start + len(questions):04d}",
            "number": number,
            "type": "judge",
            "content": content,
            "options": None,
            "answer": "",
            "analysis": "",
            "chapter": "所有章节判断题",
            "chapter_id": "ch_judge",
            "stats": {"total": 0, "correct": 0, "rate": 0},
        })
    return questions, qid_start + len(questions)


def _apply_structured_answers(
    questions: List[Dict],
    answer_data: Dict[str, Dict[str, Any]],
    judge_answers: List[bool],
) -> None:
    for question in questions:
        q_type = question.get("type")
        chapter = question.get("chapter")
        try:
            number = int(str(question.get("number") or "0"))
        except ValueError:
            continue

        if q_type == "single":
            answers = str(answer_data.get(chapter, {}).get("single", ""))
            if 1 <= number <= len(answers):
                question["answer"] = ord(answers[number - 1]) - 65
        elif q_type == "multi":
            groups = answer_data.get(chapter, {}).get("multi", [])
            if 1 <= number <= len(groups):
                question["answer"] = sorted(ord(ch) - 65 for ch in groups[number - 1] if "A" <= ch <= "D")
        elif q_type == "judge":
            if 1 <= number <= len(judge_answers):
                question["answer"] = judge_answers[number - 1]


def _validate_structured_questions(questions: List[Dict]) -> List[str]:
    issues: List[str] = []
    seen_ids: Set[str] = set()
    for question in questions:
        qid = str(question.get("id") or "")
        if qid in seen_ids:
            issues.append(f"重复题目 ID: {qid}")
        seen_ids.add(qid)

        q_type = question.get("type")
        options = question.get("options")
        answer = question.get("answer")
        label = f"{qid}/{question.get('chapter')}/{question.get('number')}"
        if q_type in {"single", "multi"}:
            if not isinstance(options, list) or len(options) != 4:
                issues.append(f"{label}: 选择题选项数不是 4")
            if answer in ("", None, []):
                issues.append(f"{label}: 缺少答案")
            if q_type == "single" and isinstance(answer, int) and not (0 <= answer < len(options or [])):
                issues.append(f"{label}: 单选答案越界")
            if q_type == "multi" and isinstance(answer, list):
                for item in answer:
                    if not isinstance(item, int) or not (0 <= item < len(options or [])):
                        issues.append(f"{label}: 多选答案越界")
        elif q_type == "judge":
            if not isinstance(answer, bool):
                issues.append(f"{label}: 判断题缺少布尔答案")
        else:
            issues.append(f"{label}: 未知题型 {q_type}")
    return issues


def _parse_structured_text_bank(text: str) -> List[Dict]:
    if "答案速查" not in text or "单项选择题" not in text:
        return []

    full_text = _strip_source_noise(text)
    answer_index = full_text.rfind("答案速查")
    if answer_index < 0:
        return []

    question_area = full_text[:answer_index]
    answer_area = full_text[answer_index:]
    answer_data, judge_answers = _parse_answer_key(answer_area)
    if not answer_data and not judge_answers:
        return []

    sections = _split_structured_sections(question_area)
    if not sections:
        return []

    questions: List[Dict] = []
    qid_next = 1
    for chapter, section_text in sections:
        if chapter == "所有章节判断题":
            judge_questions, qid_next = _parse_structured_judge_questions(section_text, qid_next)
            questions.extend(judge_questions)
            continue

        single_text, multi_text = _split_structured_choice_parts(section_text)
        if single_text:
            single_questions, qid_next = _parse_structured_choice_questions(single_text, chapter, "single", qid_next)
            questions.extend(single_questions)
        if multi_text:
            multi_questions, qid_next = _parse_structured_choice_questions(multi_text, chapter, "multi", qid_next)
            questions.extend(multi_questions)

    if not questions:
        return []

    _apply_structured_answers(questions, answer_data, judge_answers)
    issues = _validate_structured_questions(questions)
    if issues:
        sample = "；".join(issues[:12])
        raise HTTPException(status_code=400, detail=f"结构化题库解析校验失败：{sample}")

    return questions


def parse_questions_from_text(text: str) -> List[Dict]:
    """从文本解析题目 - 支持结构化题库和超紧凑格式。"""
    structured_questions = _parse_structured_text_bank(text)
    if structured_questions:
        return structured_questions

    questions = []
    full_text = text.replace('\r', '\n')

    def normalize_chapter_name(name: str) -> str:
        chapter = re.sub(r'\s+', ' ', name or "").strip()
        chapter = chapter.replace("＋", "+")
        return chapter

    # 识别章节标题（支持“绪论+第一章 ...”和“第二章 ...”等）
    chapter_pattern = re.compile(
        r'((?:绪论\s*\+\s*第[一二三四五六七八九十0-9]+章[^\n]{0,60})|(?:第[一二三四五六七八九十百0-9]+章[^\n]{0,60}))'
    )

    all_chapters = [normalize_chapter_name(m.group(1)) for m in chapter_pattern.finditer(full_text)]
    initial_chapter = next((c for c in all_chapters if "绪论" in c or "第一章" in c), "")

    chapter_to_id: Dict[str, str] = {}
    next_chapter_no = 1

    def get_chapter_id(chapter_name: str) -> str:
        nonlocal next_chapter_no
        name = normalize_chapter_name(chapter_name)
        if not name:
            return "ch01"
        if name not in chapter_to_id:
            chapter_to_id[name] = f"ch{next_chapter_no:02d}"
            next_chapter_no += 1
        return chapter_to_id[name]
    
    # 第一步：预处理文本，在题号前添加换行（方便分割）
    # 将 "...D、xxx 2、" 替换为 "...D、xxx\n2、"
    processed_text = re.sub(r'(?<=[A-Fa-f、.．\s])(\d+)([、.．])', r'\n\1\2', full_text)
    
    # 第二步：按题号分割
    raw_blocks = re.split(r'\n(?=\d+[、.．])', processed_text)
    
    qid = 1
    current_chapter = initial_chapter
    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue
        
        # 提取题号
        match = re.match(r'^(\d+)[、.．]\s*(.+)', block, re.DOTALL)
        if not match:
            continue
        
        content = match.group(2).strip()

        # 识别并剥离题干中夹带的章节标题（常见于 PDF/Word 提取）
        inline_chapter = None
        apply_to_current = False
        chapter_match = re.search(r'\n\s*' + chapter_pattern.pattern, content)
        if chapter_match:
            inline_chapter = normalize_chapter_name(chapter_match.group(1))
            content = content[:chapter_match.start()].strip()
        else:
            # 章节标题单独落在题干开头的场景
            start_match = re.match(r'^\s*' + chapter_pattern.pattern + r'\s*[\n]+', content)
            if start_match:
                inline_chapter = normalize_chapter_name(start_match.group(1))
                content = content[start_match.end():].strip()
                apply_to_current = True

        if inline_chapter and apply_to_current:
            current_chapter = inline_chapter
        
        # 初始化题目
        question = {
            "id": f"q{qid:04d}",
            "number": str(qid),
            "type": "single",
            "content": "",
            "options": [],
            "answer": "",
            "analysis": "",
            "chapter": current_chapter or "默认章节",
            "chapter_id": get_chapter_id(current_chapter or "默认章节"),
        }

        type_match = re.search(r'【(?P<type>单选题|多选题|判断题|填空题)】', content)
        if type_match:
            type_map = {"单选题": "single", "多选题": "multi", "判断题": "judge", "填空题": "blank"}
            question["type"] = type_map[type_match.group("type")]
            content = re.sub(r'【(?:单选题|多选题|判断题|填空题)】', '', content, count=1).strip()
        
        # 分离答案（在行尾或独立一行）
        answer_patterns = [
            r'答案\s*[：:]\s*([^\n]+)',
        ]
        for pattern in answer_patterns:
            answer_match = re.search(pattern, content)
            if answer_match:
                answer = answer_match.group(1).strip()
                question["answer"] = answer.upper() if re.fullmatch(r'[A-Fa-f]+', answer) else answer
                # 从内容中移除答案部分
                content = content[:answer_match.start()] + content[answer_match.end():]
                content = content.strip()
                
                if question["type"] == "blank":
                    question["answer"] = normalize_blank_answer(answer)
                elif answer in ["对", "错", "正确", "错误", "√", "×"]:
                    question["type"] = "judge"
                elif len(answer) > 1 and all(c.upper() in 'ABCDEF' for c in answer):
                    question["type"] = "multi"
                break
        
        # 分离解析
        analysis_match = re.search(r'解析\s*[：:]\s*(.+?)(?=\n\s*\d+[、.．]|\Z)', content, re.DOTALL)
        if analysis_match:
            question["analysis"] = analysis_match.group(1).strip()
            content = content[:analysis_match.start()] + content[analysis_match.end():]
            content = content.strip()
        
        # 提取选项 - 这是关键部分
        # 使用更精确的模式匹配选项
        # 匹配 A、xxx 或 A. xxx 或 A xxx 格式，直到下一个选项或行尾
        options = []
        
        # 方法1：尝试找到所有 A/B/C/D 开头的选项
        # 使用更严格的模式，确保捕获完整选项
        opt_pattern = r'([A-F])[、.．\s]+([^A-F\n]+?)(?=(?:[A-F][、.．\s]|\n|$))'
        opt_matches = list(re.finditer(opt_pattern, content, re.MULTILINE))
        
        if len(opt_matches) >= 2:
            # 找到第一个选项的位置
            first_opt_pos = opt_matches[0].start()
            question["content"] = content[:first_opt_pos].strip()
            
            for opt_match in opt_matches:
                opt_letter = opt_match.group(1)
                opt_text = opt_match.group(2).strip()
                # 清理选项文本
                opt_text = re.sub(r'\s+', ' ', opt_text)
                # 移除可能的下一题内容（如果选项文本包含数字+顿号）
                opt_text = re.sub(r'\d+[、.．].*$', '', opt_text).strip()
                options.append(opt_text)
        
        # 方法2：如果没找到足够选项，尝试按单个字母分割
        if len(options) < 2:
            # 按 A、B、C、D 分割
            parts = re.split(r'(?=[A-F])[、.．\s]*', content)
            if len(parts) > 1:
                question["content"] = parts[0].strip()
                for i, part in enumerate(parts[1:]):
                    if part.strip():
                        # 移除选项字母
                        opt_text = re.sub(r'^[A-F]\s*', '', part).strip()
                        opt_text = re.sub(r'\d+[、.．].*$', '', opt_text).strip()
                        if opt_text:
                            options.append(opt_text)
        
        question["options"] = options
        
        # 如果仍然没有选项，可能是判断题或格式错误
        if not options and question["type"] not in {"judge", "blank"}:
            question["content"] = content.strip()
            if normalize_blank_answer(question.get("answer")) and not re.fullmatch(
                r"[A-Fa-f]+",
                str(question.get("answer") or "").strip(),
            ):
                question["type"] = "blank"
                question["answer"] = normalize_blank_answer(question.get("answer"))
        
        questions.append(question)
        qid += 1

        # 章节标题落在题目末尾时，从下一题开始生效
        if inline_chapter and not apply_to_current:
            current_chapter = inline_chapter
    
    return questions
    
    return questions


@app.post("/api/extract/analyze", dependencies=[Depends(require_admin_token)])
async def generate_analysis(request: AnalyzeRequest):
    """使用 LLM 生成解析 - 高并发版本"""
    questions = request.questions
    config = request.config
    
    # 检查 LLMService 是否可用
    if LLMService is None or LLMConfig is None:
        for q in questions:
            if not q.get("analysis"):
                q["analysis"] = generate_mock_analysis(q)
        return {"questions": questions}
    
    try:
        # 支持多 API 并发（多 provider / 多 key）
        results = await _run_analysis_with_config(questions, config)
        return {"questions": results}
    
    except Exception as e:
        print(f"LLM 调用失败: {e}")
        import traceback
        traceback.print_exc()
        # 返回模拟数据作为备用
        for q in questions:
            if not q.get("analysis"):
                q["analysis"] = generate_mock_analysis(q)
        return {"questions": questions}


def generate_mock_analysis(q: Dict) -> str:
    """生成模拟解析（备用）"""
    type_names = {
        "single": "单选题",
        "multi": "多选题",
        "judge": "判断题"
    }
    
    templates = [
        f"本题考查相关知识点的理解。正确答案是 {q.get('answer', 'N/A')}。",
        f"这是一道{type_names.get(q['type'], '选择')}题，需要掌握基础概念。答案为 {q.get('answer', 'N/A')}。",
        f"解析：根据教材内容，本题正确答案为 {q.get('answer', 'N/A')}。",
    ]
    
    import random
    return random.choice(templates)


async def fill_java_answer_analyses(questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not questions:
        return questions
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=503, detail="DEEPSEEK_API_KEY 未配置，无法生成 Java 题目答案和解析")
    api_url = os.getenv("DEEPSEEK_BASE_URL", "").strip() or None
    model = os.getenv("DEEPSEEK_MODEL", "").strip() or "deepseek-chat"
    timeout = int(os.getenv("DEEPSEEK_TIMEOUT", "120") or "120")
    try:
        return await asyncio.to_thread(
            fill_java_questions_with_deepseek,
            questions,
            api_key=api_key,
            api_url=api_url,
            model=model,
            timeout=timeout,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Java 题目答案和解析生成失败: {exc}") from exc


# ============== WebSocket 实时进度 ==============

def _split_api_entries(raw: str) -> List[str]:
    """支持逗号/分号/换行分隔多个条目"""
    if not raw:
        return []
    parts = re.split(r"[,\n;]+", raw)
    return [p.strip() for p in parts if p.strip()]


def _build_config_cache_key(config: AnalysisConfig) -> str:
    payload = {
        "provider": config.provider,
        "apiKey": config.apiKey,
        "apiUrl": config.apiUrl,
        "model": config.model,
        "apiConfigs": config.apiConfigs or [],
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _parse_entry_with_fallback(entry: str, fallback: AnalysisConfig) -> "LLMConfig":
    """
    支持以下格式：
    1) 纯 Key: sk-xxx
    2) provider:key: deepseek:sk-xxx
    3) provider|key|apiUrl|model: openai|sk-xxx|https://.../v1|gpt-4o-mini
    """
    provider = fallback.provider
    api_url = fallback.apiUrl
    model = fallback.model
    api_key = entry

    if "|" in entry:
        parts = [p.strip() for p in entry.split("|")]
        if len(parts) >= 2:
            provider = parts[0] or provider
            api_key = parts[1]
        if len(parts) >= 3 and parts[2]:
            api_url = parts[2]
        if len(parts) >= 4 and parts[3]:
            model = parts[3]
    elif ":" in entry and len(entry.split(":", 1)[0]) <= 20:
        maybe_provider, maybe_key = entry.split(":", 1)
        if maybe_provider in {"openai", "deepseek", "siliconflow"} and maybe_key.strip():
            provider = maybe_provider
            api_key = maybe_key.strip()

    return LLMConfig(
        provider=provider,
        api_keys=[api_key],
        base_url=api_url,
        model=model,
        max_concurrent=5,
        timeout=30.0,
        max_retries=2
    )


def _normalize_llm_configs(config: AnalysisConfig) -> List["LLMConfig"]:
    """将前端配置归一化为多个 LLMConfig，并做缓存。"""
    cache_key = _build_config_cache_key(config)
    now = time.time()
    cached = API_CONFIG_CACHE.get(cache_key)
    if cached and now - cached[0] < API_CONFIG_CACHE_TTL:
        return cached[1]

    normalized: List["LLMConfig"] = []

    if config.apiConfigs:
        for item in config.apiConfigs:
            api_key = str(item.get("apiKey", "")).strip()
            if not api_key:
                continue
            normalized.append(
                LLMConfig(
                    provider=str(item.get("provider", config.provider or "deepseek")),
                    api_keys=[api_key],
                    base_url=item.get("apiUrl") or config.apiUrl,
                    model=item.get("model") or config.model,
                    max_concurrent=5,
                    timeout=30.0,
                    max_retries=2
                )
            )
    else:
        entries = _split_api_entries(config.apiKey)
        if not entries:
            raise ValueError("API Key 不能为空")

        plain_keys = [e for e in entries if "|" not in e and ":" not in e]
        if plain_keys and len(plain_keys) == len(entries):
            normalized.append(
                LLMConfig(
                    provider=config.provider,
                    api_keys=plain_keys,
                    base_url=config.apiUrl,
                    model=config.model,
                    max_concurrent=5,
                    timeout=30.0,
                    max_retries=2
                )
            )
        else:
            for entry in entries:
                normalized.append(_parse_entry_with_fallback(entry, config))

    if not normalized:
        raise ValueError("未解析到有效的 API 配置")

    expired = [k for k, (ts, _) in API_CONFIG_CACHE.items() if now - ts >= API_CONFIG_CACHE_TTL]
    for k in expired:
        API_CONFIG_CACHE.pop(k, None)
    API_CONFIG_CACHE[cache_key] = (now, normalized)
    return normalized


async def _run_analysis_with_config(
    questions: List[Dict],
    config: AnalysisConfig,
    progress_callback=None
) -> List[Dict]:
    configs = _normalize_llm_configs(config)

    if len(configs) == 1:
        provider = LLMService.create(configs[0])
        try:
            return await LLMService.generate_analysis_batch(
                provider,
                questions,
                progress_callback=progress_callback
            )
        finally:
            await LLMService.close_provider(provider)

    return await LLMService.generate_analysis_with_multi_keys(
        configs,
        questions,
        progress_callback=progress_callback
    )

class WebSocketProgressConfig(BaseModel):
    provider: str = "deepseek"
    apiKey: str = ""
    apiUrl: Optional[str] = None
    model: Optional[str] = None
    apiConfigs: Optional[List[Dict[str, Any]]] = None


@app.websocket("/ws/analyze/{client_id}")
async def websocket_analyze(websocket: WebSocket, client_id: str):
    """WebSocket 实时解析进度"""
    session_token = websocket.cookies.get(ADMIN_SESSION_COOKIE_NAME)
    header_token = websocket.headers.get("x-admin-token")
    if not get_admin_token() or not (
        is_admin_session_valid(session_token)
        or is_admin_token_valid(header_token)
    ):
        await websocket.close(code=1008, reason="admin session required")
        return
    await manager.connect(client_id, websocket)
    
    try:
        # 等待前端发送题目和配置
        data = await websocket.receive_json()
        questions = data.get("questions", [])
        config_data = data.get("config", {})
        
        if not questions:
            await manager.send_error(client_id, "没有题目需要解析")
            return
        
        # 检查 LLMService
        if LLMService is None or LLMConfig is None:
            # 使用模拟数据
            for i, q in enumerate(questions):
                if not q.get("analysis"):
                    q["analysis"] = generate_mock_analysis(q)
                await manager.send_progress(client_id, i + 1, len(questions), f"正在解析题目 {i+1}")
                await asyncio.sleep(0.1)  # 模拟延迟
            
            await manager.send_complete(client_id, questions)
            return
        
        try:
            config = AnalysisConfig(
                provider=config_data.get("provider", "deepseek"),
                apiKey=config_data.get("apiKey", ""),
                apiUrl=config_data.get("apiUrl"),
                model=config_data.get("model"),
                apiConfigs=config_data.get("apiConfigs")
            )
            
            # 创建进度回调
            async def progress_callback(current: int, total: int):
                await manager.send_progress(
                    client_id, 
                    current, 
                    total, 
                    f"已完成 {current}/{total} 道题目"
                )
            
            # 批量生成解析（支持多 API 异步）
            results = await _run_analysis_with_config(
                questions,
                config,
                progress_callback=progress_callback
            )
            
            await manager.send_complete(client_id, results)
            
        except Exception as e:
            print(f"WebSocket 解析失败: {e}")
            import traceback
            traceback.print_exc()
            
            # 返回模拟数据
            for q in questions:
                if not q.get("analysis"):
                    q["analysis"] = generate_mock_analysis(q)
            await manager.send_complete(client_id, questions)
    
    except WebSocketDisconnect:
        manager.disconnect(client_id)
    except Exception as e:
        await manager.send_error(client_id, str(e))
        manager.disconnect(client_id)


@app.post("/api/extract/export", dependencies=[Depends(require_admin_token)])
async def export_bank(request: ExportRequest):
    """导出标准 JSON 题库"""
    safe_name = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", request.name).strip("_") or "question_bank"
    output_file = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    output_path = os.path.join(EXPORT_DIR, output_file)
    os.makedirs(EXPORT_DIR, exist_ok=True)

    bank_data = build_standard_bank_data(request.name, request.questions)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(bank_data, f, ensure_ascii=False, indent=2)

    return {"download_url": f"/api/download/{output_file}"}


@app.post("/api/banks/save", dependencies=[Depends(require_admin_token)])
async def save_bank(request: SaveBankRequest):
    """保存题库到 tiku 目录并刷新题库列表"""
    bank_name = request.name.strip() or "未命名题库"
    bank_key = sanitize_bank_key(request.key or bank_name)
    output_path = os.path.join(TIKU_DIR, f"{bank_key}.json")

    if os.path.exists(output_path) and not request.overwrite:
        raise HTTPException(status_code=409, detail="题库代号已存在，请修改代号或启用覆盖保存")

    os.makedirs(TIKU_DIR, exist_ok=True)
    bank_data = build_standard_bank_data(bank_name, request.questions, request.color)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(bank_data, f, ensure_ascii=False, indent=2)

    saved_bank = {
        "name": bank_name,
        "color": normalize_bank_color(request.color, bank_key),
        "file": output_path,
        "data": bank_data,
    }
    sync_question_bank_to_db(bank_key, saved_bank)
    load_question_banks()

    if bank_key not in QUESTION_BANKS:
        raise HTTPException(status_code=500, detail="题库保存成功，但刷新列表失败")

    return {
        "message": "题库已保存",
        "bank": build_bank_summary(bank_key, QUESTION_BANKS[bank_key]),
        "file": os.path.relpath(output_path, BASE_DIR),
    }


@app.post("/api/banks/java/append-from-markdown", dependencies=[Depends(require_admin_token)])
async def append_java_bank_from_markdown(
    file: UploadFile = File(...),
    key: str = Form(DEFAULT_JAVA_BANK_KEY),
    start_number: int = Form(DEFAULT_JAVA_START_NUMBER),
    analyze: bool = Form(True),
    save: bool = Form(True),
):
    """从 Java Markdown 题库中增量追加新题，按题干和选项去重。"""
    if key not in QUESTION_BANKS:
        raise HTTPException(status_code=404, detail="题库不存在")

    suffix = os.path.splitext(file.filename or "")[1] or ".md"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        incoming_bank = parse_java_markdown_bank(Path(tmp_path))
    finally:
        with suppress(OSError):
            os.unlink(tmp_path)

    existing = QUESTION_BANKS[key].get("data") or {"questions": []}
    result = build_incremental_bank(
        existing,
        incoming_bank,
        start_number=start_number,
        id_prefix="java",
    )

    if analyze and result.added_questions:
        await fill_java_answer_analyses(result.added_questions)
    if save:
        output_path = os.path.join(TIKU_DIR, f"{key}.json")
        os.makedirs(TIKU_DIR, exist_ok=True)
        bank_name = result.bank.get("meta", {}).get("name") or QUESTION_BANKS[key].get("name") or "Java程序设计题库"
        bank_color = result.bank.get("meta", {}).get("color") or QUESTION_BANKS[key].get("color")
        bank_data = build_standard_bank_data(
            bank_name,
            result.bank.get("questions", []),
            bank_color,
        )
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(bank_data, f, ensure_ascii=False, indent=2)
        sync_question_bank_to_db(
            key,
            {
                "name": bank_name,
                "color": normalize_bank_color(bank_color, key),
                "file": output_path,
                "data": bank_data,
            },
        )
        load_question_banks()

    summary_bank = {
        "key": key,
        "name": result.bank.get("meta", {}).get("name") or QUESTION_BANKS[key].get("name"),
        "total": len(result.bank.get("questions", [])),
    }
    return {
        "ok": True,
        "saved": save,
        "analyzed": analyze,
        "added": result.added_count,
        "skipped_duplicates": result.skipped_duplicate_count,
        "skipped_before_start": result.skipped_before_start_count,
        "bank": summary_bank,
        "added_questions": result.added_questions,
    }


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """下载文件"""
    safe_filename = os.path.basename(filename)
    file_path = os.path.join(EXPORT_DIR, safe_filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=safe_filename)
    raise HTTPException(status_code=404, detail="文件不存在")


@app.get("/api/stats/global")
async def get_global_stats(bank: str):
    """获取题库统计"""
    require_enabled_bank(bank)

    questions = get_bank_questions(bank)
    
    stats = {
        "total": len(questions),
        "by_type": defaultdict(int),
        "by_chapter": defaultdict(int),
        "by_difficulty": {"easy": 0, "medium": 0, "hard": 0}
    }
    
    for q in questions:
        stats["by_type"][q["type"]] += 1
        stats["by_chapter"][q.get("chapter", "未知")] += 1
        
        rate = q.get("stats", {}).get("rate", 0)
        if rate >= 70:
            stats["by_difficulty"]["easy"] += 1
        elif rate >= 50:
            stats["by_difficulty"]["medium"] += 1
        else:
            stats["by_difficulty"]["hard"] += 1
    
    return {"stats": stats}


# ============== 主入口 ==============

if __name__ == "__main__":
    import uvicorn
    host = (os.getenv("APP_HOST") or "0.0.0.0").strip() or "0.0.0.0"
    raw_port = (os.getenv("APP_PORT") or os.getenv("PORT") or "10086").strip()
    try:
        port = int(raw_port)
    except ValueError:
        port = 10086

    print("🚀 启动刷题系统后端...")
    print(f"📚 API 文档: http://127.0.0.1:{port}/docs")
    uvicorn.run(app, host=host, port=port)
